"""Create a Word document mapping reported metrics to their data sources.

The script only reads static paths and writes a documentation artifact. It does
not recompute forecasts, tests, figures, or backtests.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_ROOT / "outputs_metric_source_documentation_20260524"
DOCX_PATH = OUT_DIR / "metric_data_source_documentation.docx"
MD_PATH = OUT_DIR / "metric_data_source_documentation.md"


MAINLINE_DIR = "outputs_final_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1h5_20260523"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_note_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.style = document.styles["Normal"]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)


def source_rows() -> list[dict[str, str]]:
    return [
        {
            "item": "Processed forecasting panel and sample",
            "source": "data/processed/forecasting_panel.csv",
            "scope": "Common input panel for the replication. Public-data approximation; 25 tickers; daily realized-variance forecasting panel.",
            "construction": "Contains realized variance targets and explanatory variables used to construct the model-specific design matrices. Feature standardization is fitted on the in-sample training window and then applied to validation/test features.",
            "notes": "The project uses MHAR and PARTIAL_MALL. PARTIAL_MALL is the IV-omitted version of the paper's broader M_ALL design.",
        },
        {
            "item": "Table 1 explanatory-variable summary statistics",
            "source": "outputs_paper_table1_explanatory_stats_20260523/table1_explanatory_variable_summary.csv; outputs_paper_table1_explanatory_stats_20260523/table1_explanatory_variable_summary.docx",
            "scope": "Descriptive statistics of explanatory variables from the processed panel.",
            "construction": "Reports cross-sectional means of time-series summary statistics. For asset-specific variables, bracketed ranges report the minimum and maximum of the asset-level statistic across tickers.",
            "notes": "Independent of forecast horizon and therefore unaffected by later h=22 backfills.",
        },
        {
            "item": "Main forecast prediction table",
            "source": f"{MAINLINE_DIR}/predictions/model_predictions.csv",
            "scope": "Final mainline forecast outputs for h=1 and h=5; datasets MHAR and PARTIAL_MALL; includes HAR family, regularized linear models, Bagging, RandomForest, GradientBoosting, and NN variants.",
            "construction": "Each row is a ticker-date-model-dataset-horizon forecast. The reported mainline follows the tuned-no-refit GB and NN50 outputs merged with the corrected HAR/BG/RF results.",
            "notes": "This mainline directory intentionally excludes h=22. h=22 is handled as a later separate backfill.",
        },
        {
            "item": "Forecast MSE summary by asset",
            "source": f"{MAINLINE_DIR}/tables/forecast_metrics_by_asset.csv; {MAINLINE_DIR}/tables/forecast_summary_cross_section.csv",
            "scope": "Ticker-level and cross-sectional forecast accuracy summaries for h=1 and h=5.",
            "construction": "Computes out-of-sample MSE and relative MSE by model, ticker, dataset, and horizon from the final prediction table.",
            "notes": "All forecast-evaluation numbers are test-period quantities.",
        },
        {
            "item": "Table 2-style pairwise relative MSE and Diebold-Mariano tests",
            "source": f"{MAINLINE_DIR}/tables/pairwise_relative_mse_matrix.csv; {MAINLINE_DIR}/tables/diebold_mariano_tests.csv",
            "scope": "Pairwise relative MSE matrices and ticker-level DM p-values for h=1 and h=5.",
            "construction": "Matrix cell (row i, column j) is the cross-sectional average of MSE_j / MSE_i. DM tests are computed per ticker for one-sided equal predictive accuracy tests.",
            "notes": "Table formatting follows the paper convention: more than 50 percent ticker-level rejections at 10, 5, and 1 percent are shown with italic, bold italic, and bold italic underlined formatting, respectively.",
        },
        {
            "item": "Figure 3 relative-MSE boxplots",
            "source": "outputs_figure3_relative_mse_boxplot_20260523/figure3_relative_mse_by_ticker.csv; outputs_figure3_relative_mse_boxplot_20260523/figure3_relative_mse_boxplot_paper_axis_h1.png",
            "scope": "One-day-ahead relative MSE distributions across tickers for MHAR and PARTIAL_MALL.",
            "construction": "Uses ticker-level relative MSE values derived from forecast_metrics_by_asset.csv. Boxplots summarize cross-sectional dispersion by model.",
            "notes": "The paper-axis version fixes the y-axis to match the visual scale of the original figure; the underlying CSV contains the plotted values.",
        },
        {
            "item": "Figure 4 MCS inclusion rates",
            "source": "outputs_mcs_hln_closest_h1_20260523/tables/mcs_inclusion_rates.csv; outputs_figure4_mcs_hln_closest_h1_20260523/figure4_mcs_inclusion_rate_h1.png",
            "scope": "One-day-ahead Model Confidence Set inclusion rates for MHAR and PARTIAL_MALL.",
            "construction": "Runs an HLN-style MCS procedure on model loss series and reports the percentage of times each model is retained in the 90 percent confidence set.",
            "notes": "Use this as an appendix robustness figure unless the report has space for an MCS-focused discussion.",
        },
        {
            "item": "Figure 5 forecast accuracy over realized-volatility deciles",
            "source": "outputs_rv_decile_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1h5_20260523/tables/rv_decile_mse.csv; outputs_figure5_rv_decile_paper_style_20260523/figure5_rv_decile_paper_style_h1.png",
            "scope": "One-day-ahead PARTIAL_MALL results for selected representative models.",
            "construction": "Splits the test sample into realized-variance deciles and reports each selected model's MSE relative to HAR within the selected decile groups.",
            "notes": "Models follow the paper-style representative subset: HAR-X, LogHAR, EN, RF, and NN2_10.",
        },
        {
            "item": "Figure 6 ALE between explanatory variables and future volatility",
            "source": "outputs_figure6_ale_paper_formula_20260524/tables/ale_table_paper_formula.csv; outputs_figure6_ale_paper_formula_20260524/figures/figure6_ale_paper_formula_raw_autoaxis_h1.png",
            "scope": "AAPL, PARTIAL_MALL, h=1; features RVD, RVW, and M1W; representative models HAR-X, LogHAR, EN, RF, and NN2_10.",
            "construction": "Computes centered ALE following the paper's finite-difference formula using the training information set Z. The plotted y-values are centered ALE estimates.",
            "notes": "The display-scaled figure multiplies ALE by 10^3 for readability; this is a display convention and does not change the underlying ALE table.",
        },
        {
            "item": "Figure 7 variable-importance measure",
            "source": "outputs_variable_importance_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1_20260523/tables/variable_importance.csv; outputs_figure7_variable_importance_paper_style_h1_20260524/figure7_variable_importance_plot_data.csv",
            "scope": "PARTIAL_MALL, h=1, 25 tickers; representative models HAR-X, EN, RF, and NN2_10.",
            "construction": "Computes ALE-based variable importance as the standard deviation of centered ALE curves, normalized so each model's importances sum to one, then averages across tickers.",
            "notes": "The current plotted source predates the binary-EA ALE fix. Recompute VI later if the final Figure 7 should show a small nonzero EA bar.",
        },
        {
            "item": "VaR summary table",
            "source": "outputs_var_summary_table_h1_20260524/var_summary_table_wide.csv; outputs_var_summary_table_h1_20260524/var_summary_table_h1.docx",
            "scope": "FHS VaR application, h=1, MHAR and PARTIAL_MALL.",
            "construction": "Reports relative VaR quantile/check loss versus HAR, the share of ticker-level loss tests favoring the selected model, exceedance rates, and unconditional/independence coverage rejection rates.",
            "notes": "This is a compact application table; it is easier to use in the main text than the full pairwise VaR matrix.",
        },
        {
            "item": "Table 8/9-style pairwise relative VaR and DM tests",
            "source": "outputs_pairwise_var_dm_tables_h1_20260524/pairwise_relative_var_loss_matrix.csv; outputs_pairwise_var_dm_tables_h1_20260524/var_diebold_mariano_tests.csv; outputs_pairwise_var_dm_tables_h1_20260524/var_bottom_rows.csv; outputs_pairwise_var_dm_tables_h1_20260524/pairwise_relative_var_dm_tables_h1.docx",
            "scope": "FHS VaR application, h=1, datasets MHAR and PARTIAL_MALL.",
            "construction": "Matrix cell (row i, column j) is the cross-sectional average of VaR check-loss_j / check-loss_i. Bottom rows report exceedance probability, unconditional coverage rejections, and conditional coverage rejections.",
            "notes": "The second table is PARTIAL_MALL rather than the paper's full M_ALL because IV is omitted in the available replication dataset.",
        },
        {
            "item": "Table A.6 hyperparameter settings",
            "source": "table_a6_hyperparameters_reproduction.docx; config/paper_core_rolling_tuned_no_refit.yaml; config/paper_core_rolling.yaml; src/rv1rep/nn.py",
            "scope": "Mainline model-tuning and training settings.",
            "construction": "Documents regularization grids, tree settings, GB 40-grid, NN learning rate, epochs, patience, initializer, ensemble size, and 50-seed procedure.",
            "notes": "Mainline NN uses Keras Dropout(rate=0.8). A separate robustness check interprets the reported 0.8 as a retention probability and reruns selected NN cases with Dropout(rate=0.2).",
        },
        {
            "item": "Corrected dropout/y-standardization robustness run",
            "source": "outputs_corrected_paper_style_h1h5_20260524_143417_active/",
            "scope": "Separate active robustness experiment for selected tickers at h=1 and h=5.",
            "construction": "Runs corrected NN variants using the 50-seed procedure and revised target-scaling/dropout interpretation. This output is isolated from all mainline tables.",
            "notes": "Exclude from final tables until the run is complete and the generated CSVs have passed validation.",
        },
    ]


def write_markdown(rows: list[dict[str, str]]) -> None:
    lines = [
        "# Data Sources and Construction of Reported Metrics",
        "",
        f"All paths are relative to the project root. The mainline forecast results are taken from `{MAINLINE_DIR}/` unless a row states otherwise.",
        "",
        "The table below documents the source files, sample scope, and construction rule for each reported table or figure. It is intended as a copyable source map for the report or appendix.",
        "",
        "| Reported item | Primary data source | Scope | Construction | Notes |",
        "|---|---|---|---|---|",
    ]
    for row in rows:
        lines.append(
            "| "
            + " | ".join(
                row[key].replace("|", "\\|").replace("\n", " ")
                for key in ["item", "source", "scope", "construction", "notes"]
            )
            + " |"
        )
    lines.extend(
        [
            "",
            "## Interpretation Notes",
            "",
            "- All forecast-accuracy, DM, MCS, and VaR evaluation metrics are computed on the out-of-sample test period.",
            "- The current mainline documentation covers h=1 and h=5. It does not include later h=22 backfills.",
            "- PARTIAL_MALL should be described as the IV-omitted version of the broader all-covariate design, not as the paper's full M_ALL.",
            "- VaR tables use VaR check loss for relative-loss and DM calculations, not raw VaR levels.",
            "- Existing Figure 7 source files predate the binary-feature ALE fix for EA; rerun the VI script later if a nonzero EA bar is required.",
            "- The active corrected dropout/y-standardization robustness run is documented only as an isolated experiment until it finishes and passes validation.",
            "",
        ]
    )
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx(rows: list[dict[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Data Sources and Construction of Reported Metrics")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    add_note_paragraph(
        document,
        "All paths are relative to the project root. The mainline forecast results are taken from "
        f"{MAINLINE_DIR}/ unless a row states otherwise.",
    )
    add_note_paragraph(
        document,
        "This document is a source map for the reported tables and figures. It records the input files, "
        "sample scope, and construction rule for each metric without recomputing any results.",
    )

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = [
        "Reported item / metric",
        "Primary data source",
        "Scope",
        "Construction",
        "Notes / caveats",
    ]
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, size=8)

    widths = [Cm(3.5), Cm(6.8), Cm(4.8), Cm(6.6), Cm(5.8)]
    for row in rows:
        cells = table.add_row().cells
        values = [row["item"], row["source"], row["scope"], row["construction"], row["notes"]]
        for cell, value, width in zip(cells, values, widths):
            cell.width = width
            set_cell_text(cell, value, size=7)

    document.add_paragraph()
    p = document.add_paragraph()
    r = p.add_run("Interpretation notes")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    notes = [
        "All forecast-accuracy, DM, MCS, and VaR evaluation metrics are computed on the out-of-sample test period.",
        "The current mainline documentation covers h=1 and h=5. It does not include later h=22 backfills.",
        "PARTIAL_MALL should be described as the IV-omitted version of the broader all-covariate design, not as the paper's full M_ALL.",
        "VaR tables use VaR check loss for relative-loss and DM calculations, not raw VaR levels.",
        "Existing Figure 7 source files predate the binary-feature ALE fix for EA; rerun the VI script later if a nonzero EA bar is required.",
        "The active corrected dropout/y-standardization robustness run is documented only as an isolated experiment until it finishes and passes validation.",
    ]
    for note in notes:
        p = document.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("- " + note)
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)

    document.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = source_rows()
    write_markdown(rows)
    build_docx(rows)
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
