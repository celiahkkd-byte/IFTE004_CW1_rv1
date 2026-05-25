from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


MODEL_ORDER = [
    'HAR',
    'HARX',
    'LogHAR',
    'LevHAR',
    'SHAR',
    'HARQ',
    'Ridge',
    'Lasso',
    'ElasticNet',
    'AdaptiveLasso',
    'PostLasso',
    'Bagging',
    'RandomForest',
    'GradientBoosting',
    'NN1_1',
    'NN1',
    'NN2_1',
    'NN2',
    'NN3_1',
    'NN3',
    'NN4_1',
    'NN4',
]

PLAIN_LABELS = {
    'HAR': 'HAR',
    'HARX': 'HAR-X',
    'LogHAR': 'LogHAR',
    'LevHAR': 'LevHAR',
    'SHAR': 'SHAR',
    'HARQ': 'HARQ',
    'Ridge': 'RR',
    'Lasso': 'LA',
    'ElasticNet': 'EN',
    'AdaptiveLasso': 'A-LA',
    'PostLasso': 'P-LA',
    'Bagging': 'BG',
    'RandomForest': 'RF',
    'GradientBoosting': 'GB',
}

NN_LABELS = {
    'NN1_1': (1, 1),
    'NN1': (10, 1),
    'NN2_1': (1, 2),
    'NN2': (10, 2),
    'NN3_1': (1, 3),
    'NN3': (10, 3),
    'NN4_1': (1, 4),
    'NN4': (10, 4),
}

HORIZON_LABELS = {
    1: 'One-day-ahead',
    5: 'One-week-ahead',
    22: 'One-month-ahead',
}

DATASET_SUBSCRIPTS = {
    'MHAR': 'HAR',
    'PARTIAL_MALL': 'PARTIAL_MALL',
}


def resolve(path: str | Path) -> Path:
    p = Path(path)
    return p.resolve()


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.28)
    section.bottom_margin = Inches(0.28)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(8)


def set_cell_borders(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = f'w:{edge}'
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f'w:{key}'), str(value))


def clear_borders(cell) -> None:
    nil = {'val': 'nil'}
    set_cell_borders(cell, top=nil, left=nil, bottom=nil, right=nil, insideH=nil, insideV=nil)


def horizontal_rule(cell, top: bool = False, bottom: bool = False) -> None:
    attrs = {}
    if top:
        attrs['top'] = {'val': 'single', 'sz': '6', 'space': '0', 'color': '666666'}
    if bottom:
        attrs['bottom'] = {'val': 'single', 'sz': '6', 'space': '0', 'color': '666666'}
    if attrs:
        set_cell_borders(cell, **attrs)


def set_cell_width(cell, width: float) -> None:
    cell.width = Inches(width)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in('w:tcW')
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(int(width * 1440)))
    tc_w.set(qn('w:type'), 'dxa')


def set_run_font(run, size: float, bold: bool = False, italic: bool = False, underline: bool = False) -> None:
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def add_model_label(paragraph, model: str, size: float = 7.2, bold: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if model in NN_LABELS:
        power, subscript = NN_LABELS[model]
        run = paragraph.add_run('NN')
        set_run_font(run, size, bold=bold)
        sup = paragraph.add_run(str(power))
        set_run_font(sup, size - 1.2, bold=bold)
        sup.font.superscript = True
        sub = paragraph.add_run(str(subscript))
        set_run_font(sub, size - 1.2, bold=bold)
        sub.font.subscript = True
        return
    run = paragraph.add_run(PLAIN_LABELS.get(model, model))
    set_run_font(run, size, bold=bold)


def add_dataset_symbol(paragraph, dataset: str, size: float = 10.5) -> None:
    run = paragraph.add_run('M')
    set_run_font(run, size)
    sub = paragraph.add_run(DATASET_SUBSCRIPTS.get(dataset, dataset))
    set_run_font(sub, size - 1)
    sub.font.subscript = True
    if dataset == 'PARTIAL_MALL':
        note = paragraph.add_run(' (IV omitted)')
        set_run_font(note, size)


def add_caption(doc: Document, table_number: int, dataset: str, horizon: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    first = p.add_run(f'Table {table_number} ')
    set_run_font(first, 10.5, bold=True)
    rest = p.add_run(f'{HORIZON_LABELS.get(horizon, f"h={horizon}")} relative MSE and Diebold-Mariano test for dataset ')
    set_run_font(rest, 10.5)
    add_dataset_symbol(p, dataset, size=10.5)


def add_note(doc: Document, extra_note: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Notes: ')
    set_run_font(r, 8.2, italic=True)
    note = (
        'We report the out-of-sample realized variance forecast MSE of each model in the selected column '
        'relative to the benchmark in the selected row. Each number is a cross-sectional average of such '
        'pairwise relative MSEs for each stock. Formatting is as follows: italic, bold italic, and bold '
        'italic underlined denote that the Diebold-Mariano test of equal predictive accuracy is rejected '
        'for more than 50% of stocks at the 10%, 5%, and 1% significance levels, respectively. The hypothesis '
        'being tested is H0: MSE_i = MSE_j against the one-sided alternative H1: MSE_i > MSE_j, where model i '
        'is the selected row and model j is the selected column.'
    )
    r = p.add_run(note)
    set_run_font(r, 8.2)
    if extra_note:
        r = p.add_run(' ' + extra_note)
        set_run_font(r, 8.2)


def significance_style(dm: pd.DataFrame, dataset: str, horizon: int, row_model: str, col_model: str) -> tuple[str, float, float, float]:
    if row_model == col_model:
        return 'none', 0.0, 0.0, 0.0
    g = dm[
        (dm['dataset'] == dataset)
        & (dm['horizon'] == horizon)
        & (dm['row_model'] == row_model)
        & (dm['col_model'] == col_model)
    ]
    if g.empty:
        raise ValueError(f'Missing DM tests for dataset={dataset} h={horizon} row={row_model} col={col_model}')
    p = pd.to_numeric(g['p_value'], errors='coerce').dropna()
    if p.empty:
        return 'none', 0.0, 0.0, 0.0
    share10 = float((p < 0.10).mean())
    share5 = float((p < 0.05).mean())
    share1 = float((p < 0.01).mean())
    if share1 > 0.5:
        return 'p01', share10, share5, share1
    if share5 > 0.5:
        return 'p05', share10, share5, share1
    if share10 > 0.5:
        return 'p10', share10, share5, share1
    return 'none', share10, share5, share1


def add_value(paragraph, text: str, style: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bold = style in {'p05', 'p01'}
    italic = style in {'p10', 'p05', 'p01'}
    underline = style == 'p01'
    run = paragraph.add_run(text)
    set_run_font(run, 7.2, bold=bold, italic=italic, underline=underline)


def build_table(doc: Document, matrix: pd.DataFrame, dm: pd.DataFrame, dataset: str, horizon: int) -> pd.DataFrame:
    subset = matrix[(matrix['dataset'] == dataset) & (matrix['horizon'] == horizon)].copy()
    if subset.empty:
        raise ValueError(f'Missing pairwise matrix for dataset={dataset} horizon={horizon}')
    subset = subset.set_index('benchmark_row')

    missing_rows = [m for m in MODEL_ORDER if m not in subset.index]
    missing_cols = [m for m in MODEL_ORDER if m not in subset.columns]
    if missing_rows or missing_cols:
        raise ValueError(f'Missing models in matrix rows={missing_rows} cols={missing_cols}')

    table = doc.add_table(rows=len(MODEL_ORDER) + 1, cols=len(MODEL_ORDER) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    row_label_width = 0.62
    numeric_width = 0.475
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            clear_borders(cell)
            set_cell_width(cell, row_label_width if idx == 0 else numeric_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

    for cell in table.rows[0].cells:
        horizontal_rule(cell, top=True, bottom=True)
    for cell in table.rows[-1].cells:
        horizontal_rule(cell, bottom=True)

    table.cell(0, 0).text = ''
    for j, model in enumerate(MODEL_ORDER, start=1):
        p = table.cell(0, j).paragraphs[0]
        add_model_label(p, model, size=7.2, bold=False)

    audit_rows = []
    for i, row_model in enumerate(MODEL_ORDER, start=1):
        p = table.cell(i, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_model_label(p, row_model, size=7.2, bold=False)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for j, col_model in enumerate(MODEL_ORDER, start=1):
            p = table.cell(i, j).paragraphs[0]
            if row_model == col_model:
                add_value(p, '-', 'none')
                audit_rows.append(
                    {
                        'dataset': dataset,
                        'horizon': horizon,
                        'row_model': row_model,
                        'col_model': col_model,
                        'reject_share_10pct': 0.0,
                        'reject_share_5pct': 0.0,
                        'reject_share_1pct': 0.0,
                        'applied_style': 'none',
                    }
                )
                continue
            value = float(subset.loc[row_model, col_model])
            style, share10, share5, share1 = significance_style(dm, dataset, horizon, row_model, col_model)
            add_value(p, f'{value:.3f}', style)
            audit_rows.append(
                {
                    'dataset': dataset,
                    'horizon': horizon,
                    'row_model': row_model,
                    'col_model': col_model,
                    'reject_share_10pct': share10,
                    'reject_share_5pct': share5,
                    'reject_share_1pct': share1,
                    'applied_style': style,
                }
            )
    return pd.DataFrame(audit_rows)


def load_inputs(tables_dir: Path) -> tuple[pd.DataFrame, pd.DataFrame]:
    matrix_path = tables_dir / 'pairwise_relative_mse_matrix.csv'
    dm_path = tables_dir / 'diebold_mariano_tests.csv'
    if not matrix_path.exists():
        raise FileNotFoundError(f'Missing pairwise matrix: {matrix_path}')
    if not dm_path.exists():
        raise FileNotFoundError(f'Missing DM tests: {dm_path}')
    matrix = pd.read_csv(matrix_path)
    dm = pd.read_csv(dm_path)
    matrix['horizon'] = matrix['horizon'].astype(int)
    dm['horizon'] = dm['horizon'].astype(int)
    return matrix, dm


def main() -> None:
    ap = argparse.ArgumentParser(
        description='Create Word tables for pairwise relative MSE and Diebold-Mariano formatting.'
    )
    ap.add_argument('--tables-dir', required=True)
    ap.add_argument('--output-dir', required=True)
    ap.add_argument('--datasets', nargs='+', default=['MHAR', 'PARTIAL_MALL'])
    ap.add_argument('--horizons', nargs='+', type=int, default=[1, 5])
    ap.add_argument('--output-name', default='pairwise_relative_mse_dm_tables_h1h5.docx')
    ap.add_argument('--table-number-start', type=int, default=2)
    ap.add_argument('--extra-note', default=None)
    args = ap.parse_args()

    tables_dir = resolve(args.tables_dir)
    output_dir = resolve(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    matrix, dm = load_inputs(tables_dir)

    doc = Document()
    set_doc_defaults(doc)
    audit_parts = []

    table_number = args.table_number_start
    first = True
    for horizon in args.horizons:
        for dataset in args.datasets:
            if not first:
                doc.add_page_break()
            first = False
            add_caption(doc, table_number, dataset, horizon)
            audit_parts.append(build_table(doc, matrix, dm, dataset, horizon))
            add_note(doc, extra_note=args.extra_note)
            table_number += 1

    output_path = output_dir / args.output_name
    doc.save(output_path)
    audit = pd.concat(audit_parts, ignore_index=True)
    audit_path = output_dir / output_path.with_suffix('.formatting_audit.csv').name
    audit.to_csv(audit_path, index=False)
    print(f'Wrote {output_path}')
    print(f'Wrote {audit_path}')


if __name__ == '__main__':
    main()
