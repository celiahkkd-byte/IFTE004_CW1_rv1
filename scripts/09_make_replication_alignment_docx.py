"""Create a Word appendix comparing the mainline replication with the paper.

This document is intentionally scoped to the reported mainline forecast results:

outputs_final_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1h5_20260523/

It does not document later robustness experiments.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs_replication_alignment_20260524"
DOCX_PATH = OUT_DIR / "replication_alignment_original_paper.docx"
MD_PATH = OUT_DIR / "replication_alignment_original_paper.md"

MAINLINE_DIR = "outputs_final_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1h5_20260523"
MAINLINE_PRED = f"{MAINLINE_DIR}/predictions/model_predictions.csv"
PAPER = "Christensen, Siggaard, and Veliyev (2023)"


def alignment_rows() -> list[dict[str, str]]:
    return [
        {
            "component": "Main result scope",
            "paper": "The paper reports one-day, one-week, and one-month realized-variance forecasts for M_HAR and M_ALL.",
            "replication": (
                f"The reported mainline is {MAINLINE_DIR}/. It contains h=1 and h=5 only, "
                "for MHAR and PARTIAL_MALL. It has 1,831,652 prediction rows, 25 tickers, "
                "22 model labels, and no duplicated ticker/date/dataset/horizon/model keys."
            ),
            "evidence": f"{MAINLINE_PRED}; run_provenance.json",
            "alignment": "Partial alignment. h=1 and h=5 are covered; h=22 is deliberately excluded from this mainline.",
        },
        {
            "component": "Data source and assets",
            "paper": (
                "The paper uses cleaned NYSE TAQ transaction data for 29 DJIA constituents, "
                "with predecessor histories where needed, over 2001-2017."
            ),
            "replication": (
                "The mainline uses public Alpha Vantage five-minute OHLCV files for 25 tickers "
                "with available continuous coverage. It does not reconstruct unavailable predecessor histories."
            ),
            "evidence": "rv1.pdf, Section 2; config/paper_core_rolling_tuned_no_refit.yaml: assets.tickers and raw_dir",
            "alignment": "Deviation caused by public-data availability. Model design is replicated, but the raw data source and cross-section are not identical.",
        },
        {
            "component": "Realized variance construction",
            "paper": "RV is the sum of squared intraday log returns. The paper uses 5-minute returns with n=78.",
            "replication": (
                "compute_daily_realized_measures forms 5-minute log returns and computes "
                "rv = sum(r5**2) * rv_scale. In the mainline config, rv_scale = 1.0."
            ),
            "evidence": "rv1.pdf, Eq. (2) and Section 2; src/rv1rep/intraday.py; config assets.rv_scale",
            "alignment": "Aligned. No annualization or extra scaling is applied in model inputs or targets.",
        },
        {
            "component": "Forecast target",
            "paper": (
                "For longer horizons, the target is the average realized variance from t+1 to t+h, "
                "conditional on information available at t."
            ),
            "replication": (
                "h=1 uses s.shift(-1). For h>1, future_average uses "
                "s.shift(-1).rolling(h).mean().shift(-(h-1)), which aligns each date t with "
                "mean(RV_{t+1},...,RV_{t+h})."
            ),
            "evidence": "rv1.pdf, Section 4; src/rv1rep/features.py::_future_target",
            "alignment": "Aligned for h=1 and h=5.",
        },
        {
            "component": "Feature sets",
            "paper": (
                "M_HAR contains RVD, RVW, and RVM. M_ALL extends this set with IV, EA, M1W, "
                "$VOL, VIX, HSI, ADS, US3M, and EPU."
            ),
            "replication": (
                "MHAR is implemented with HAR lags and model-specific HAR extensions. "
                "The extended dataset is named PARTIAL_MALL because public IV is unavailable in the mainline. "
                "The code includes IV only if an IV file exists; otherwise the remaining public predictors are used."
            ),
            "evidence": "rv1.pdf, Table 1 and Section 2; src/rv1rep/features.py::feature_columns_for_model",
            "alignment": "MHAR aligned. Extended dataset is a partial M_ALL replication because IV is omitted.",
        },
        {
            "component": "M1W construction",
            "paper": "The paper defines M1W as one-week momentum. Formula (8) belongs to the LevHAR aggregated return variables, not to M1W.",
            "replication": "m1w is the five-trading-day cumulative close-to-close log return ending at t.",
            "evidence": "rv1.pdf, Section 2 and Eq. (8); src/rv1rep/features.py::add_asset_features",
            "alignment": "Aligned with the standard momentum interpretation and with the scale of the paper's Table 1.",
        },
        {
            "component": "Transformed public variables",
            "paper": "US3M is first differenced, $VOL is first log-differenced, and LogHAR log-transforms VIX and IV.",
            "replication": (
                "us3m_diff is constructed by first differencing the 3-month Treasury bill series. "
                "dvol is the first difference of log dollar volume. LogHAR uses log_vix and log_iv when available."
            ),
            "evidence": "rv1.pdf, Section 2; src/rv1rep/features.py::merge_external and feature_columns_for_model",
            "alignment": "Aligned, except IV-related transformations are unavailable when IV is absent.",
        },
        {
            "component": "Feature standardization",
            "paper": "The paper standardizes input data using the training-set sample mean and variance before estimation.",
            "replication": (
                "The mainline run standardizes all feature columns supplied to each model within the in-sample training window, "
                "including the EA dummy when it appears in PARTIAL_MALL. The target remains in realized-variance units."
            ),
            "evidence": "rv1.pdf, Section 2 and Appendix A.3 note; config preprocessing.standardize_binary_features=true; src/rv1rep/preprocessing.py",
            "alignment": "Aligned for feature standardization. Target standardization is not part of the reported mainline.",
        },
        {
            "component": "Train/validation/test split",
            "paper": "The primary split is 70% training, 10% validation, and 20% test, chronologically ordered.",
            "replication": "The mainline config uses train_frac=0.70, val_frac=0.10, and test_frac=0.20.",
            "evidence": "rv1.pdf, Section 2; config/paper_core_rolling_tuned_no_refit.yaml: splitting",
            "alignment": "Aligned.",
        },
        {
            "component": "HAR family estimation",
            "paper": "Non-regularized HAR models merge training and validation and use rolling-window forecasts.",
            "replication": (
                "HAR, HARX, LogHAR, LevHAR, SHAR, and HARQ use the combined train+validation rolling window. "
                "The output records n_val=0 for these models because no validation tuning is performed."
            ),
            "evidence": "rv1.pdf, Section 2; src/rv1rep/forecasting.py::_fit_one_asset_rolling; src/rv1rep/models.py::fit_sklearn_model",
            "alignment": "Aligned.",
        },
        {
            "component": "Regularized linear estimation",
            "paper": "RR, LA, EN, A-LA, and P-LA tune hyperparameters on validation data under a rolling scheme without train-validation concatenation.",
            "replication": (
                "Ridge, Lasso, ElasticNet, AdaptiveLasso, and PostLasso are rolling daily-refit models. "
                "Candidate models fit on the training block, validation selects hyperparameters, and the selected estimator is not refit on train+validation."
            ),
            "evidence": (
                "rv1.pdf, Section 2 and Appendix A.4; src/rv1rep/models.py::_maybe_refit_selected_tuned_model; "
                "mainline params include fit_sample='train_only_after_validation_selection'"
            ),
            "alignment": "Design aligned. Hyperparameter grid density differs from the paper; see model table.",
        },
        {
            "component": "Bagging and Random Forest estimation",
            "paper": "BG and RF adopt the rolling train+validation approach and use default tree settings.",
            "replication": (
                "BG and RF use the combined train+validation window and default-like settings. "
                "For computational tractability, the mainline refits these models every five test observations rather than every test observation."
            ),
            "evidence": "rv1.pdf, Section 2 and Table A.6; config estimation.ml_refit_every=5; src/rv1rep/forecasting.py",
            "alignment": "Partially aligned. Window construction and hyperparameters align, but refit frequency is an implementation deviation.",
        },
        {
            "component": "Gradient Boosting estimation",
            "paper": "GB tunes depth, number of trees, and learning rate on validation data under the rolling no-concatenation design.",
            "replication": (
                "GB uses depth in {1,2}, trees 50 to 500, learning rates {0.01,0.1}, validation-tuned no-refit. "
                "Like the tree models, it is refitted every five test observations in the mainline."
            ),
            "evidence": "rv1.pdf, Section 2 and Table A.6; config gradient_boosting; src/rv1rep/models.py; run_provenance.json",
            "alignment": "Grid and no-refit design aligned; refit frequency is a computational deviation.",
        },
        {
            "component": "Neural-network estimation",
            "paper": (
                "NNs use fixed-window estimation, architectures NN1-NN4, Leaky ReLU, Adam, learning rate 0.001, "
                "dropout 0.8, patience 100, epochs 500, Glorot normal initialization, and best/top-10 selection from 100 seeds."
            ),
            "replication": (
                "The mainline keeps fixed-window NN forecasts and the same architecture and listed hyperparameters. "
                "It uses 50 seeds, with single-best and top-10 ensemble forecasts selected by validation MSE."
            ),
            "evidence": "rv1.pdf, Sections 1.5, 1.6, 2, Appendix A.5; config neural_network; outputs_nn50 checkpoint params in mainline predictions",
            "alignment": "Partially aligned. Architecture and training design align; seed count is 50 rather than the paper's 100.",
        },
        {
            "component": "LogHAR back-transformation",
            "paper": "LogHAR forecasts log realized variance and applies a Jensen-type bias correction when converting back to RV.",
            "replication": "LogHAR computes exp(predicted log RV + 0.5 * residual variance) before evaluation.",
            "evidence": "rv1.pdf, footnote near Eq. (26); src/rv1rep/forecasting.py; src/rv1rep/models.py::FittedModel.predict",
            "alignment": "Aligned.",
        },
        {
            "component": "Forecast evaluation",
            "paper": "Forecasts are evaluated on the test period using out-of-sample MSE, pairwise relative MSE, and one-sided Diebold-Mariano tests.",
            "replication": (
                "The mainline produces forecast_summary_cross_section.csv, forecast_metrics_by_asset.csv, "
                "pairwise_relative_mse_matrix.csv, and diebold_mariano_tests.csv from the final test-period predictions."
            ),
            "evidence": "rv1.pdf, Section 1.6 and Tables 2-7; src/rv1rep/evaluation.py; mainline tables directory",
            "alignment": "Aligned for the h=1 and h=5 mainline tables.",
        },
    ]


def model_rows() -> list[dict[str, str]]:
    return [
        {
            "paper_model": "HAR",
            "rep_model": "HAR",
            "implementation": "OLS on RVD, RVW, RVM; rolling train+validation window.",
            "alignment": "Aligned.",
        },
        {
            "paper_model": "HAR-X",
            "rep_model": "HARX",
            "implementation": "OLS HAR with extended public predictors in PARTIAL_MALL; IV omitted when unavailable.",
            "alignment": "Aligned in MHAR; partial for M_ALL because IV is missing.",
        },
        {
            "paper_model": "LogHAR",
            "rep_model": "LogHAR",
            "implementation": "Log-transformed HAR variables and log target; Jensen correction back to RV.",
            "alignment": "Aligned; extended IV term unavailable in PARTIAL_MALL.",
        },
        {
            "paper_model": "LevHAR",
            "rep_model": "LevHAR",
            "implementation": "Adds daily, weekly, and monthly negative return terms; weekly/monthly returns are averages as in Eq. (8).",
            "alignment": "Aligned.",
        },
        {
            "paper_model": "SHAR",
            "rep_model": "SHAR",
            "implementation": "Uses positive and negative realized semivariance plus RVW and RVM.",
            "alignment": "Aligned.",
        },
        {
            "paper_model": "HARQ",
            "rep_model": "HARQ",
            "implementation": "Uses sqrt(RQ_t) x RVD_t interaction plus HAR lags.",
            "alignment": "Aligned.",
        },
        {
            "paper_model": "RR",
            "rep_model": "Ridge",
            "implementation": "Ridge regression; validation-tuned lambda; no train+validation refit.",
            "alignment": "Design aligned; grid is 80 log-spaced lambda values rather than the paper's 1000.",
        },
        {
            "paper_model": "LA",
            "rep_model": "Lasso",
            "implementation": "Lasso; validation-tuned lambda; no train+validation refit.",
            "alignment": "Design aligned; grid is 80 log-spaced lambda values rather than the paper's 1000.",
        },
        {
            "paper_model": "EN",
            "rep_model": "ElasticNet",
            "implementation": "Elastic net; lambda grid and alpha/l1-ratio grid selected by validation.",
            "alignment": "Design aligned; grid is 80 lambda values and seven alpha values, smaller than the paper's 1000 x 10 grid.",
        },
        {
            "paper_model": "A-LA",
            "rep_model": "AdaptiveLasso",
            "implementation": "Two-step weighted lasso using OLS-based adaptive weights; validation-tuned lambda.",
            "alignment": "Conceptually aligned; exact numerical implementation may differ from the paper's software.",
        },
        {
            "paper_model": "P-LA",
            "rep_model": "PostLasso",
            "implementation": "First-stage lasso selects variables; second-stage OLS fits selected columns on the training block.",
            "alignment": "Aligned in design; lambda grid smaller than the paper's.",
        },
        {
            "paper_model": "BG",
            "rep_model": "Bagging",
            "implementation": "BaggingRegressor with 500 trees and min_samples_leaf=5; combined train+validation window.",
            "alignment": "Hyperparameters aligned; refit every five test observations rather than daily.",
        },
        {
            "paper_model": "RF",
            "rep_model": "RandomForest",
            "implementation": "RandomForestRegressor with 500 trees, min_samples_leaf=5, and max_features=floor(J/3).",
            "alignment": "Hyperparameters aligned; refit every five test observations rather than daily.",
        },
        {
            "paper_model": "GB",
            "rep_model": "GradientBoosting",
            "implementation": "Validation grid: depth {1,2}, trees {50,100,...,500}, learning rate {0.01,0.1}; no-refit.",
            "alignment": "Grid aligned; Python scikit-learn implementation and five-observation refit frequency differ from the paper's R implementation/daily rolling interpretation.",
        },
        {
            "paper_model": "NN1, NN2, NN3, NN4; single-best and top-10 ensemble",
            "rep_model": "NN1_1/NN1, NN2_1/NN2, NN3_1/NN3, NN4_1/NN4",
            "implementation": "Architectures [2], [4,2], [8,4,2], [16,8,4,2]; LeakyReLU, Adam, dropout=0.8, epochs=500, patience=100.",
            "alignment": "Architecture and hyperparameters aligned; 50 seeds instead of 100.",
        },
    ]


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_para(doc: Document, text: str, *, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def build_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[Cm]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, size=8)
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            cell.width = width
            set_cell_text(cell, value, size=7)
    doc.add_paragraph()


def write_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Replication Alignment with the Original Paper")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    add_para(
        doc,
        f"Scope: this alignment table covers only the reported mainline forecast results in {MAINLINE_DIR}/. "
        "It excludes later dropout or target-standardization robustness experiments.",
    )
    add_para(
        doc,
        f"Primary prediction table: {MAINLINE_PRED}. The mainline contains h=1 and h=5, MHAR and PARTIAL_MALL, "
        "GradientBoosting, and NN50 forecasts.",
    )

    add_para(doc, "Table 1. Design and estimation alignment", bold=True, size=11)
    build_table(
        doc,
        ["Component", f"Original paper: {PAPER}", "Mainline replication", "Evidence checked", "Alignment"],
        [[r["component"], r["paper"], r["replication"], r["evidence"], r["alignment"]] for r in alignment_rows()],
        [Cm(3.2), Cm(6.0), Cm(7.0), Cm(5.4), Cm(5.4)],
    )

    add_para(doc, "Table 2. Model-by-model mapping", bold=True, size=11)
    build_table(
        doc,
        ["Original paper label", "Mainline label", "Mainline implementation", "Alignment"],
        [[r["paper_model"], r["rep_model"], r["implementation"], r["alignment"]] for r in model_rows()],
        [Cm(5.0), Cm(4.0), Cm(10.0), Cm(8.0)],
    )

    add_para(
        doc,
        "Summary: the mainline replication is closest to the paper in the realized-variance construction, target alignment, "
        "chronological split, HAR-family rolling design, validation-tuned no-refit regularized/GB design, NN architecture, "
        "and MSE/DM evaluation. The main deviations are the public-data cross-section, PARTIAL_MALL instead of full M_ALL "
        "because IV is unavailable, h=22 exclusion, 50 NN seeds rather than 100, smaller regularization grids, and the "
        "five-observation refit interval for BG/RF/GB.",
        bold=False,
    )

    doc.save(DOCX_PATH)


def write_markdown() -> None:
    lines: list[str] = []
    lines.append("# Replication Alignment with the Original Paper")
    lines.append("")
    lines.append(
        f"Scope: this alignment table covers only `{MAINLINE_DIR}/` and excludes later robustness experiments."
    )
    lines.append(f"Primary prediction table: `{MAINLINE_PRED}`.")
    lines.append("")
    lines.append("## Table 1. Design and estimation alignment")
    lines.append("")
    lines.append("| Component | Original paper | Mainline replication | Evidence checked | Alignment |")
    lines.append("|---|---|---|---|---|")
    for r in alignment_rows():
        lines.append(
            f"| {r['component']} | {r['paper']} | {r['replication']} | {r['evidence']} | {r['alignment']} |"
        )
    lines.append("")
    lines.append("## Table 2. Model-by-model mapping")
    lines.append("")
    lines.append("| Original paper label | Mainline label | Mainline implementation | Alignment |")
    lines.append("|---|---|---|---|")
    for r in model_rows():
        lines.append(f"| {r['paper_model']} | {r['rep_model']} | {r['implementation']} | {r['alignment']} |")
    lines.append("")
    lines.append(
        "Summary: the mainline replication is closest to the paper in realized-variance construction, "
        "target alignment, chronological split, HAR-family rolling design, validation-tuned no-refit "
        "regularized/GB design, NN architecture, and MSE/DM evaluation. The main deviations are the "
        "public-data cross-section, PARTIAL_MALL instead of full M_ALL because IV is unavailable, h=22 "
        "exclusion, 50 NN seeds rather than 100, smaller regularization grids, and the five-observation "
        "refit interval for BG/RF/GB."
    )
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    write_docx()
    write_markdown()
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
