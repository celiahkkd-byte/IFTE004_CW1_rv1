from __future__ import annotations

import argparse
from pathlib import Path
import sys

import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from rv1rep.evaluation import diebold_mariano


DEFAULT_VAR_FORECASTS = (
    ROOT
    / "outputs_fhs_var_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1_20260523"
    / "predictions"
    / "var_forecasts_fhs.csv"
)
DEFAULT_VAR_SUMMARY = (
    ROOT
    / "outputs_fhs_var_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1_20260523"
    / "tables"
    / "var_backtest_fhs_summary.csv"
)
DEFAULT_OUTPUT_DIR = ROOT / "outputs_pairwise_var_dm_tables_h1_20260524"

MODEL_ORDER = [
    "HAR",
    "HARX",
    "LogHAR",
    "LevHAR",
    "SHAR",
    "HARQ",
    "Ridge",
    "Lasso",
    "ElasticNet",
    "AdaptiveLasso",
    "PostLasso",
    "Bagging",
    "RandomForest",
    "GradientBoosting",
    "NN1_1",
    "NN1",
    "NN2_1",
    "NN2",
    "NN3_1",
    "NN3",
    "NN4_1",
    "NN4",
]

PLAIN_LABELS = {
    "HAR": "HAR",
    "HARX": "HAR-X",
    "LogHAR": "LogHAR",
    "LevHAR": "LevHAR",
    "SHAR": "SHAR",
    "HARQ": "HARQ",
    "Ridge": "RR",
    "Lasso": "LA",
    "ElasticNet": "EN",
    "AdaptiveLasso": "A-LA",
    "PostLasso": "P-LA",
    "Bagging": "BG",
    "RandomForest": "RF",
    "GradientBoosting": "GB",
}

NN_LABELS = {
    "NN1_1": (1, 1),
    "NN1": (10, 1),
    "NN2_1": (1, 2),
    "NN2": (10, 2),
    "NN3_1": (1, 3),
    "NN3": (10, 3),
    "NN4_1": (1, 4),
    "NN4": (10, 4),
}

DATASET_SUBSCRIPTS = {
    "MHAR": "HAR",
    "PARTIAL_MALL": "PARTIAL_MALL",
}


def resolve(path: str | Path) -> Path:
    p = Path(path)
    if not p.is_absolute():
        p = ROOT / p
    return p.resolve()


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.28)
    section.bottom_margin = Inches(0.28)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(8)


def set_run_font(run, size: float, bold: bool = False, italic: bool = False, underline: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def set_cell_borders(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_borders(cell) -> None:
    nil = {"val": "nil"}
    set_cell_borders(cell, top=nil, left=nil, bottom=nil, right=nil, insideH=nil, insideV=nil)


def horizontal_rule(cell, top: bool = False, bottom: bool = False) -> None:
    attrs = {}
    if top:
        attrs["top"] = {"val": "single", "sz": "6", "space": "0", "color": "666666"}
    if bottom:
        attrs["bottom"] = {"val": "single", "sz": "6", "space": "0", "color": "666666"}
    if attrs:
        set_cell_borders(cell, **attrs)


def set_cell_width(cell, width: float) -> None:
    cell.width = Inches(width)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_model_label(paragraph, model: str, size: float = 7.2, bold: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if model in NN_LABELS:
        power, subscript = NN_LABELS[model]
        run = paragraph.add_run("NN")
        set_run_font(run, size, bold=bold)
        sup = paragraph.add_run(str(power))
        set_run_font(sup, size - 1.2, bold=bold)
        sup.font.superscript = True
        sub = paragraph.add_run(str(subscript))
        set_run_font(sub, size - 1.2, bold=bold)
        sub.font.subscript = True
        return
    run = paragraph.add_run(PLAIN_LABELS.get(model, model))
    set_run_font(run, size, bold=bold)


def add_dataset_symbol(paragraph, dataset: str, size: float = 10.5) -> None:
    run = paragraph.add_run("M")
    set_run_font(run, size)
    sub = paragraph.add_run(DATASET_SUBSCRIPTS.get(dataset, dataset))
    set_run_font(sub, size - 1)
    sub.font.subscript = True
    if dataset == "PARTIAL_MALL":
        note = paragraph.add_run(" (IV omitted)")
        set_run_font(note, size)


def add_caption(doc: Document, table_number: int, dataset: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    first = p.add_run(f"Table {table_number} ")
    set_run_font(first, 10.5, bold=True)
    rest = p.add_run("One-day-ahead relative VaR check loss and Diebold-Mariano test for dataset ")
    set_run_font(rest, 10.5)
    add_dataset_symbol(p, dataset, size=10.5)


def add_note(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Notes: ")
    set_run_font(r, 8.2, italic=True)
    note = (
        "We report one-day-ahead 5% filtered-historical-simulation VaR check loss of each model in "
        "the selected column relative to the benchmark in the selected row. Each number is a "
        "cross-sectional average of ticker-level pairwise relative VaR check losses. Formatting is "
        "as follows: italic, bold italic, and bold italic underlined denote that the one-sided "
        "Diebold-Mariano test of equal VaR check loss is rejected for more than 50% of stocks at the "
        "10%, 5%, and 1% significance levels, respectively. The hypothesis is H0: E(L_i)=E(L_j) "
        "against H1: E(L_i)>E(L_j), where model i is the selected row and model j is the selected "
        "column. Prb. is the average exceedance rate. Unc. is the number of stocks for which the "
        "Kupiec unconditional coverage test rejects at the 5% level. Cond. is the number of stocks "
        "for which the Christoffersen conditional coverage test, computed as LR_uc + LR_ind with "
        "two degrees of freedom, rejects at the 5% level."
    )
    r = p.add_run(note)
    set_run_font(r, 8.2)


def add_value(paragraph, text: str, style: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bold = style in {"p05", "p01"}
    italic = style in {"p10", "p05", "p01"}
    underline = style == "p01"
    run = paragraph.add_run(text)
    set_run_font(run, 7.2, bold=bold, italic=italic, underline=underline)


def significance_style(dm: pd.DataFrame, dataset: str, row_model: str, col_model: str) -> tuple[str, float, float, float]:
    if row_model == col_model:
        return "none", 0.0, 0.0, 0.0
    g = dm[(dm["dataset"] == dataset) & (dm["row_model"] == row_model) & (dm["col_model"] == col_model)]
    if g.empty:
        raise ValueError(f"Missing VaR DM tests dataset={dataset} row={row_model} col={col_model}")
    p = pd.to_numeric(g["p_value"], errors="coerce").dropna()
    if p.empty:
        return "none", 0.0, 0.0, 0.0
    share10 = float((p < 0.10).mean())
    share5 = float((p < 0.05).mean())
    share1 = float((p < 0.01).mean())
    if share1 > 0.5:
        return "p01", share10, share5, share1
    if share5 > 0.5:
        return "p05", share10, share5, share1
    if share10 > 0.5:
        return "p10", share10, share5, share1
    return "none", share10, share5, share1


def load_var_forecasts(path: Path, datasets: list[str]) -> pd.DataFrame:
    usecols = ["date", "ticker", "dataset", "horizon", "model", "var_loss"]
    df = pd.read_csv(path, usecols=usecols, parse_dates=["date"])
    df = df[
        df["dataset"].isin(datasets)
        & df["horizon"].astype(int).eq(1)
        & df["model"].astype(str).isin(MODEL_ORDER)
    ].copy()
    df["ticker"] = df["ticker"].astype(str).str.upper()
    df["model"] = df["model"].astype(str)
    df["var_loss"] = pd.to_numeric(df["var_loss"], errors="coerce")
    df = df.dropna(subset=["date", "ticker", "dataset", "model", "var_loss"])
    if df.empty:
        raise ValueError("No VaR forecast rows remain after filtering.")
    return df


def compute_pairwise_var_loss(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    matrix_rows = []
    dm_rows = []
    for dataset, gd in df.groupby("dataset", sort=True):
        for row_model in MODEL_ORDER:
            row = {"dataset": dataset, "benchmark_row": row_model}
            for col_model in MODEL_ORDER:
                ratios = []
                for ticker, gt in gd.groupby("ticker", sort=True):
                    a = gt[gt["model"].eq(row_model)][["date", "var_loss"]].rename(columns={"var_loss": "loss_i"})
                    b = gt[gt["model"].eq(col_model)][["date", "var_loss"]].rename(columns={"var_loss": "loss_j"})
                    merged = a.merge(b, on="date", how="inner")
                    if merged.empty:
                        continue
                    loss_i = merged["loss_i"].to_numpy(dtype=float)
                    loss_j = merged["loss_j"].to_numpy(dtype=float)
                    mean_i = float(np.mean(loss_i))
                    mean_j = float(np.mean(loss_j))
                    if mean_i > 0:
                        ratios.append(mean_j / mean_i)
                    stat, p = diebold_mariano(loss_i, loss_j, alternative="greater", h=1)
                    dm_rows.append(
                        {
                            "dataset": dataset,
                            "ticker": ticker,
                            "row_model": row_model,
                            "col_model": col_model,
                            "dm_stat": stat,
                            "p_value": p,
                            "n": int(len(merged)),
                        }
                    )
                row[col_model] = float(np.mean(ratios)) if ratios else np.nan
            matrix_rows.append(row)
    return pd.DataFrame(matrix_rows), pd.DataFrame(dm_rows)


def load_bottom_rows(path: Path, datasets: list[str]) -> pd.DataFrame:
    summary = pd.read_csv(path)
    required = {
        "dataset",
        "horizon",
        "ticker",
        "model",
        "exceedance_rate",
        "kupiec_p",
        "kupiec_lr",
        "christoffersen_ind_lr",
    }
    missing = required - set(summary.columns)
    if missing:
        raise ValueError(f"Missing columns in VaR summary: {sorted(missing)}")
    summary = summary[
        summary["dataset"].isin(datasets)
        & summary["horizon"].astype(int).eq(1)
        & summary["model"].astype(str).isin(MODEL_ORDER)
    ].copy()
    summary["cond_lr"] = pd.to_numeric(summary["kupiec_lr"], errors="coerce") + pd.to_numeric(
        summary["christoffersen_ind_lr"], errors="coerce"
    )
    summary["cond_p"] = 1.0 - stats.chi2.cdf(summary["cond_lr"], 2)
    return (
        summary.groupby(["dataset", "model"], as_index=False)
        .agg(
            prob=("exceedance_rate", "mean"),
            unc=("kupiec_p", lambda s: int((pd.to_numeric(s, errors="coerce") < 0.05).sum())),
            cond=("cond_p", lambda s: int((pd.to_numeric(s, errors="coerce") < 0.05).sum())),
            n_tickers=("ticker", "nunique"),
        )
        .reset_index(drop=True)
    )


def add_bottom_value(paragraph, text: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, 7.2)


def build_table(
    doc: Document,
    matrix: pd.DataFrame,
    dm: pd.DataFrame,
    bottom: pd.DataFrame,
    dataset: str,
) -> pd.DataFrame:
    subset = matrix[matrix["dataset"].eq(dataset)].copy().set_index("benchmark_row")
    if subset.empty:
        raise ValueError(f"Missing pairwise VaR matrix for dataset={dataset}")
    missing_rows = [m for m in MODEL_ORDER if m not in subset.index]
    missing_cols = [m for m in MODEL_ORDER if m not in subset.columns]
    if missing_rows or missing_cols:
        raise ValueError(f"Missing models in VaR matrix rows={missing_rows} cols={missing_cols}")

    table = doc.add_table(rows=len(MODEL_ORDER) + 4, cols=len(MODEL_ORDER) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    row_label_width = 0.62
    numeric_width = 0.475
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            clear_borders(cell)
            set_cell_width(cell, row_label_width if idx == 0 else numeric_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

    for cell in table.rows[0].cells:
        horizontal_rule(cell, top=True, bottom=True)
    for cell in table.rows[len(MODEL_ORDER)].cells:
        horizontal_rule(cell, bottom=True)
    for cell in table.rows[-1].cells:
        horizontal_rule(cell, bottom=True)

    table.cell(0, 0).text = ""
    for j, model in enumerate(MODEL_ORDER, start=1):
        add_model_label(table.cell(0, j).paragraphs[0], model, size=7.2)

    audit_rows = []
    for i, row_model in enumerate(MODEL_ORDER, start=1):
        p = table.cell(i, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_model_label(p, row_model, size=7.2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for j, col_model in enumerate(MODEL_ORDER, start=1):
            p = table.cell(i, j).paragraphs[0]
            if row_model == col_model:
                add_value(p, "-", "none")
                audit_rows.append(
                    {
                        "dataset": dataset,
                        "row_model": row_model,
                        "col_model": col_model,
                        "reject_share_10pct": 0.0,
                        "reject_share_5pct": 0.0,
                        "reject_share_1pct": 0.0,
                        "applied_style": "none",
                    }
                )
                continue
            value = float(subset.loc[row_model, col_model])
            style, share10, share5, share1 = significance_style(dm, dataset, row_model, col_model)
            add_value(p, f"{value:.3f}", style)
            audit_rows.append(
                {
                    "dataset": dataset,
                    "row_model": row_model,
                    "col_model": col_model,
                    "reject_share_10pct": share10,
                    "reject_share_5pct": share5,
                    "reject_share_1pct": share1,
                    "applied_style": style,
                }
            )

    bottom_index = bottom[bottom["dataset"].eq(dataset)].set_index("model")
    labels = [("Prb.", "prob"), ("Unc.", "unc"), ("Cond.", "cond")]
    for offset, (label, col) in enumerate(labels, start=len(MODEL_ORDER) + 1):
        p = table.cell(offset, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_bottom_value(p, label)
        for j, model in enumerate(MODEL_ORDER, start=1):
            value = bottom_index.loc[model, col]
            text = f"{float(value):.3f}" if col == "prob" else str(int(value))
            add_bottom_value(table.cell(offset, j).paragraphs[0], text)
    return pd.DataFrame(audit_rows)


def main() -> None:
    parser = argparse.ArgumentParser(description="Create paper-style pairwise relative VaR check-loss/DM Word tables.")
    parser.add_argument("--var-forecasts", default=str(DEFAULT_VAR_FORECASTS))
    parser.add_argument("--var-summary", default=str(DEFAULT_VAR_SUMMARY))
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR))
    parser.add_argument("--datasets", nargs="+", default=["MHAR", "PARTIAL_MALL"])
    parser.add_argument("--table-number-start", type=int, default=8)
    parser.add_argument("--output-name", default="pairwise_relative_var_dm_tables_h1.docx")
    parser.add_argument("--allow-existing-output-dir", action="store_true")
    args = parser.parse_args()

    output_dir = resolve(args.output_dir)
    if output_dir.exists() and not args.allow_existing_output_dir:
        files = [p for p in output_dir.rglob("*") if p.is_file()]
        if files:
            raise SystemExit(
                f"Output directory already contains files: {output_dir}. "
                "Use a fresh directory or pass --allow-existing-output-dir."
            )
    output_dir.mkdir(parents=True, exist_ok=True)

    var_df = load_var_forecasts(resolve(args.var_forecasts), args.datasets)
    matrix, dm = compute_pairwise_var_loss(var_df)
    bottom = load_bottom_rows(resolve(args.var_summary), args.datasets)

    matrix_path = output_dir / "pairwise_relative_var_loss_matrix.csv"
    dm_path = output_dir / "var_diebold_mariano_tests.csv"
    bottom_path = output_dir / "var_bottom_rows.csv"
    matrix.to_csv(matrix_path, index=False)
    dm.to_csv(dm_path, index=False)
    bottom.to_csv(bottom_path, index=False)

    doc = Document()
    set_doc_defaults(doc)
    audit_parts = []
    table_number = int(args.table_number_start)
    first = True
    for dataset in args.datasets:
        if not first:
            doc.add_page_break()
        first = False
        add_caption(doc, table_number, dataset)
        audit_parts.append(build_table(doc, matrix, dm, bottom, dataset))
        add_note(doc)
        table_number += 1

    output_path = output_dir / args.output_name
    doc.save(output_path)
    audit = pd.concat(audit_parts, ignore_index=True)
    audit_path = output_dir / "pairwise_relative_var_dm_formatting_audit.csv"
    audit.to_csv(audit_path, index=False)

    print(f"Wrote {output_path}")
    print(f"Wrote {matrix_path}")
    print(f"Wrote {dm_path}")
    print(f"Wrote {bottom_path}")
    print(f"Wrote {audit_path}")


if __name__ == "__main__":
    main()
