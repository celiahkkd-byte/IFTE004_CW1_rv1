"""Create a Word document describing variable-level data sources.

This is a documentation-only script. It reads the processed panel to report
coverage counts, but it does not rebuild the panel or recompute any forecasts.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_ROOT / "outputs_variable_source_documentation_20260524"
DOCX_PATH = OUT_DIR / "variable_data_source_documentation.docx"
MD_PATH = OUT_DIR / "variable_data_source_documentation.md"
PANEL_PATH = PROJECT_ROOT / "data/processed/forecasting_panel.csv"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 7) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_paragraph(document: Document, text: str, *, bold: bool = False, size: int = 10) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def source_inventory() -> list[dict[str, str]]:
    return [
        {
            "source": "Intraday stock prices and volume",
            "files": "data/external/alpha_vantage_intraday_5min/combined_teacher_format/<TICKER>.txt",
            "provider": "Alpha Vantage five-minute OHLCV bars",
            "used_for": "RV, RVP, RVN, RQ, close-to-close returns, momentum, and dollar-volume change",
            "notes": "The original paper uses cleaned TAQ transaction prices. This replication uses public Alpha Vantage five-minute OHLCV bars as a transparent substitute.",
        },
        {
            "source": "Earnings announcements",
            "files": "data/external/earnings_announcements.csv; data/external/alpha_vantage_earnings/",
            "provider": "Alpha Vantage EARNINGS endpoint",
            "used_for": "EA binary indicator",
            "notes": "Reported dates are converted to firm-day indicators; non-trading reported dates are aligned to the next available trading day.",
        },
        {
            "source": "Market implied volatility",
            "files": "data/external/vix_fred.csv",
            "provider": "FRED VIXCLS",
            "used_for": "VIX and log(VIX)",
            "notes": "Merged by calendar date and forward-filled to the equity trading calendar.",
        },
        {
            "source": "Three-month Treasury bill rate",
            "files": "data/external/us3m_fred.csv",
            "provider": "FRED DTB3",
            "used_for": "US3M, stored as us3m_diff",
            "notes": "The project uses the first difference of the 3-month T-bill rate, matching the nonstationarity treatment in the paper.",
        },
        {
            "source": "Economic Policy Uncertainty",
            "files": "data/external/epu_daily.csv",
            "provider": "policyuncertainty.com daily US policy index",
            "used_for": "EPU",
            "notes": "Daily level series, forward-filled to the equity trading calendar.",
        },
        {
            "source": "ADS business conditions index",
            "files": "data/external/ads.csv",
            "provider": "Federal Reserve Bank of Philadelphia ADS current vintage",
            "used_for": "ADS",
            "notes": "Daily/near-daily level series, forward-filled to the equity trading calendar.",
        },
        {
            "source": "Hang Seng Index",
            "files": "data/external/hsi.csv",
            "provider": "Yahoo Finance symbol ^HSI",
            "used_for": "HSI",
            "notes": "Constructed as the squared daily log return of the Hang Seng closing price.",
        },
        {
            "source": "Firm-level implied volatility",
            "files": "data/external/firm_iv.csv",
            "provider": "User-supplied, paper source is OptionMetrics",
            "used_for": "IV and log(IV), if supplied",
            "notes": "This file is absent in the current project. Therefore the main extended dataset is PARTIAL_MALL, which omits IV.",
        },
    ]


def variable_dictionary() -> list[dict[str, str]]:
    return [
        {
            "acronym": "RV target",
            "panel_cols": "rv; target_rv_h1; target_rv_h5; target_log_rv_h1; target_log_rv_h5",
            "source": "Alpha Vantage five-minute intraday OHLCV bars",
            "construction": "RV_t is the sum of 78 five-minute squared log returns over the regular trading day. h=1 target is RV_{t+1}; h=5 target is the future average RV over t+1,...,t+5.",
            "scope": "Target variable; all datasets",
            "notes": "Log targets are used only by LogHAR and are transformed back for forecast evaluation.",
        },
        {
            "acronym": "RVD",
            "panel_cols": "rvd",
            "source": "Derived from daily realized variance",
            "construction": "Current-day realized variance RV_t.",
            "scope": "Asset-specific; MHAR and PARTIAL_MALL",
            "notes": "In Table 1 only, reported as annualized volatility percent: sqrt(252 x RV) x 100.",
        },
        {
            "acronym": "RVW",
            "panel_cols": "rvw",
            "source": "Derived from daily realized variance",
            "construction": "Five-trading-day rolling mean of RV ending at t.",
            "scope": "Asset-specific; MHAR and PARTIAL_MALL",
            "notes": "In Table 1 only, reported as annualized volatility percent.",
        },
        {
            "acronym": "RVM",
            "panel_cols": "rvm",
            "source": "Derived from daily realized variance",
            "construction": "Twenty-two-trading-day rolling mean of RV ending at t.",
            "scope": "Asset-specific; MHAR and PARTIAL_MALL",
            "notes": "In Table 1 only, reported as annualized volatility percent.",
        },
        {
            "acronym": "RVP / RVN",
            "panel_cols": "rvp; rvn",
            "source": "Derived from intraday five-minute returns",
            "construction": "Positive and negative realized semivariances: sums of squared positive and negative five-minute returns.",
            "scope": "Asset-specific; SHAR model",
            "notes": "These are model-specific HAR extension variables rather than Table 1 explanatory-variable rows.",
        },
        {
            "acronym": "RQ / HARQ term",
            "panel_cols": "rq; sqrt_rq_x_rvd",
            "source": "Derived from intraday five-minute returns",
            "construction": "RQ is realized quarticity. The HARQ predictor is sqrt(RQ_t) x RVD_t.",
            "scope": "Asset-specific; HARQ model",
            "notes": "Used to allow realized-volatility persistence to vary with measurement noise.",
        },
        {
            "acronym": "Leverage terms",
            "panel_cols": "rd; rw; rm",
            "source": "Derived from daily close-to-close log returns",
            "construction": "RD is min(0, daily return); RW and RM are the negative parts of the five-day and twenty-two-day rolling mean close-to-close returns.",
            "scope": "Asset-specific; LevHAR model",
            "notes": "These terms capture asymmetric volatility responses to negative returns.",
        },
        {
            "acronym": "IV",
            "panel_cols": "iv; log_iv",
            "source": "Optional data/external/firm_iv.csv; paper source is OptionMetrics",
            "construction": "If provided, firm-level implied volatility is merged by date and ticker. LogHAR uses log_iv.",
            "scope": "Asset-specific; PARTIAL_MALL only if supplied",
            "notes": "Not available in the current project, so main extended results are IV-omitted PARTIAL_MALL.",
        },
        {
            "acronym": "EA",
            "panel_cols": "ea",
            "source": "Alpha Vantage EARNINGS endpoint",
            "construction": "Binary firm-day indicator equal to one on an earnings-announcement date and zero otherwise.",
            "scope": "Asset-specific; PARTIAL_MALL",
            "notes": "Categorical 0/1 variable. Descriptive moments are omitted in Table 1, following the paper's convention. In the reported mainline estimation, EA was standardized with the other feature columns when included; later robustness code treats EA as a non-standardized categorical feature.",
        },
        {
            "acronym": "VIX",
            "panel_cols": "vix; log_vix",
            "source": "FRED VIXCLS",
            "construction": "Daily VIX level merged by date and forward-filled on the equity trading calendar. LogHAR uses log_vix.",
            "scope": "Market-wide; PARTIAL_MALL",
            "notes": "Represents market-level option-implied volatility.",
        },
        {
            "acronym": "EPU",
            "panel_cols": "epu",
            "source": "policyuncertainty.com daily US Economic Policy Uncertainty index",
            "construction": "Daily policy-uncertainty index level merged by date and forward-filled to the equity trading calendar.",
            "scope": "Macro-wide; PARTIAL_MALL",
            "notes": "Used in levels in the forecasting panel.",
        },
        {
            "acronym": "US3M",
            "panel_cols": "us3m_diff",
            "source": "FRED DTB3",
            "construction": "First difference of the three-month Treasury bill rate after sorting by date and forward-filling the level series.",
            "scope": "Macro-wide; PARTIAL_MALL",
            "notes": "The project variable is US3M, not US1M. Table 1 multiplies this transformed variable by 100.",
        },
        {
            "acronym": "HSI",
            "panel_cols": "hsi",
            "source": "Yahoo Finance ^HSI daily close",
            "construction": "Squared daily log return of the Hang Seng Index closing price.",
            "scope": "Market-wide; PARTIAL_MALL",
            "notes": "Table 1 multiplies this transformed variable by 100.",
        },
        {
            "acronym": "M1W",
            "panel_cols": "m1w",
            "source": "Derived from each stock's daily close-to-close log returns",
            "construction": "One-week momentum: five-trading-day rolling sum of close-to-close log returns ending at t.",
            "scope": "Asset-specific; PARTIAL_MALL",
            "notes": "Table 1 multiplies this transformed variable by 100.",
        },
        {
            "acronym": "$VOL",
            "panel_cols": "dvol",
            "source": "Derived from Alpha Vantage intraday close prices and volume",
            "construction": "Dollar volume is the intraday sum of close x volume. $VOL is the first difference of log dollar volume.",
            "scope": "Asset-specific; PARTIAL_MALL",
            "notes": "Stored as dvol in the panel. Table 1 multiplies dvol by 100.",
        },
        {
            "acronym": "ADS",
            "panel_cols": "ads",
            "source": "Federal Reserve Bank of Philadelphia ADS index",
            "construction": "ADS business conditions index level merged by date and forward-filled to the equity trading calendar.",
            "scope": "Macro-wide; PARTIAL_MALL",
            "notes": "Table 1 multiplies this transformed variable by 100.",
        },
    ]


def add_coverage(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    if not PANEL_PATH.exists():
        for row in rows:
            row["coverage"] = "Panel file not found"
        return rows
    panel = pd.read_csv(PANEL_PATH, parse_dates=["date"])
    total = len(panel)
    tickers = panel["ticker"].nunique() if "ticker" in panel.columns else 0
    date_min = panel["date"].min().date() if "date" in panel.columns and not panel.empty else "n/a"
    date_max = panel["date"].max().date() if "date" in panel.columns and not panel.empty else "n/a"
    for row in rows:
        cols = [c.strip() for c in row["panel_cols"].replace(";", ",").split(",")]
        cols = [c for c in cols if c in panel.columns]
        if not cols:
            row["coverage"] = "Not present in current panel"
            continue
        nonmissing = panel[cols].notna().all(axis=1).sum()
        row["coverage"] = f"{nonmissing:,}/{total:,} rows nonmissing; {tickers} tickers; {date_min} to {date_max}"
    return rows


def write_markdown(inventory: list[dict[str, str]], variables: list[dict[str, str]]) -> None:
    lines = [
        "# Variable Data Sources and Transformations",
        "",
        "This document describes the data source and construction of the variables used in the replication. Paths are relative to the project root.",
        "",
        "## Source Inventory",
        "",
        "| Source | Project file(s) | Provider | Used for | Notes |",
        "|---|---|---|---|---|",
    ]
    for row in inventory:
        lines.append("| " + " | ".join(row[k] for k in ["source", "files", "provider", "used_for", "notes"]) + " |")
    lines.extend(
        [
            "",
            "## Variable Dictionary",
            "",
            "| Variable | Panel column(s) | Source | Construction | Scope | Coverage | Notes |",
            "|---|---|---|---|---|---|---|",
        ]
    )
    for row in variables:
        lines.append(
            "| "
            + " | ".join(row[k].replace("|", "\\|") for k in ["acronym", "panel_cols", "source", "construction", "scope", "coverage", "notes"])
            + " |"
        )
    lines.extend(
        [
            "",
            "## Key Replication Notes",
            "",
            "- The current project uses a public-data replication panel. Intraday realized measures are based on Alpha Vantage five-minute OHLCV bars, whereas the paper uses cleaned TAQ transaction data.",
            "- Firm-level OptionMetrics IV is not available in the current project. Results labelled PARTIAL_MALL are therefore IV-omitted extended-model results.",
            "- Macro and market variables are merged by date and forward-filled on the equity trading calendar. This uses previously available observations and avoids look-ahead.",
            "- Table 1 applies display transformations for several variables. These display transformations do not imply that the model-estimation panel stores those variables in the same units.",
            "- In the reported mainline results, feature standardization is fitted within each in-sample training window and applied to validation/test features. This mainline standardization is applied to all feature columns supplied to the model, including EA when present in PARTIAL_MALL. Later robustness code treats EA as a categorical feature and excludes it from standardization.",
            "",
        ]
    )
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[Cm], font_size: int = 7) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for cell, text, width in zip(cells, row, widths):
            cell.width = width
            set_cell_text(cell, text, size=font_size)


def build_docx(inventory: list[dict[str, str]], variables: list[dict[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Variable Data Sources and Transformations")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)

    add_paragraph(
        document,
        "This document records the data source, construction, and project column mapping for the variables used in the replication. "
        "All paths are relative to the project root.",
    )
    add_paragraph(
        document,
        "The final modelling panel is data/processed/forecasting_panel.csv. The extended public-data design is labelled PARTIAL_MALL because firm-level IV is omitted unless data/external/firm_iv.csv is supplied.",
    )

    add_paragraph(document, "Source inventory", bold=True, size=12)
    add_table(
        document,
        ["Source", "Project file(s)", "Provider", "Used for", "Notes"],
        [[row[k] for k in ["source", "files", "provider", "used_for", "notes"]] for row in inventory],
        [Cm(3.3), Cm(6.2), Cm(4.2), Cm(4.8), Cm(8.0)],
        font_size=7,
    )

    document.add_paragraph()
    add_paragraph(document, "Variable dictionary", bold=True, size=12)
    add_table(
        document,
        ["Variable", "Panel column(s)", "Source", "Construction", "Scope", "Coverage", "Notes"],
        [
            [row[k] for k in ["acronym", "panel_cols", "source", "construction", "scope", "coverage", "notes"]]
            for row in variables
        ],
        [Cm(2.1), Cm(3.4), Cm(4.1), Cm(7.0), Cm(3.3), Cm(4.0), Cm(4.9)],
        font_size=6,
    )

    document.add_paragraph()
    add_paragraph(document, "Key replication notes", bold=True, size=12)
    notes = [
        "The current project uses a public-data replication panel. Intraday realized measures are based on Alpha Vantage five-minute OHLCV bars, whereas the paper uses cleaned TAQ transaction data.",
        "Firm-level OptionMetrics IV is not available in the current project. Results labelled PARTIAL_MALL are therefore IV-omitted extended-model results.",
        "Macro and market variables are merged by date and forward-filled on the equity trading calendar. This uses previously available observations and avoids look-ahead.",
        "Table 1 applies display transformations for several variables. These display transformations do not imply that the model-estimation panel stores those variables in the same units.",
        "In the reported mainline results, feature standardization is fitted within each in-sample training window and applied to validation/test features. This mainline standardization is applied to all feature columns supplied to the model, including EA when present in PARTIAL_MALL. Later robustness code treats EA as a categorical feature and excludes it from standardization.",
    ]
    for note in notes:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("- " + note)
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)

    document.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    inventory = source_inventory()
    variables = add_coverage(variable_dictionary())
    write_markdown(inventory, variables)
    build_docx(inventory, variables)
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
