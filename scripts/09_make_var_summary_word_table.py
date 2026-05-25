from __future__ import annotations

import argparse
import json
import os
import sys
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from rv1rep.evaluation import diebold_mariano

DEFAULT_INPUT = (
    ROOT
    / "outputs_fhs_var_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1_20260523"
    / "tables"
    / "var_backtest_fhs_summary.csv"
)
DEFAULT_VAR_FORECASTS = (
    ROOT
    / "outputs_fhs_var_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1_20260523"
    / "predictions"
    / "var_forecasts_fhs.csv"
)
DEFAULT_OUTPUT_DIR = ROOT / "outputs_var_summary_table_h1_20260524"

MODEL_ORDER = [
    "HAR",
    "HARX",
    "LogHAR",
    "LevHAR",
    "SHAR",
    "HARQ",
    "Ridge",
    "Lasso",
    "ElasticNet",
    "AdaptiveLasso",
    "PostLasso",
    "Bagging",
    "RandomForest",
    "GradientBoosting",
    "NN1_1",
    "NN1",
    "NN2_1",
    "NN2",
    "NN3_1",
    "NN3",
    "NN4_1",
    "NN4",
]

PLAIN_LABELS = {
    "HAR": "HAR",
    "HARX": "HAR-X",
    "LogHAR": "LogHAR",
    "LevHAR": "LevHAR",
    "SHAR": "SHAR",
    "HARQ": "HARQ",
    "Ridge": "RR",
    "Lasso": "LA",
    "ElasticNet": "EN",
    "AdaptiveLasso": "A-LA",
    "PostLasso": "P-LA",
    "Bagging": "BG",
    "RandomForest": "RF",
    "GradientBoosting": "GB",
}

NN_LABELS = {
    "NN1_1": (1, 1),
    "NN1": (10, 1),
    "NN2_1": (1, 2),
    "NN2": (10, 2),
    "NN3_1": (1, 3),
    "NN3": (10, 3),
    "NN4_1": (1, 4),
    "NN4": (10, 4),
}

DATASET_ORDER = ["MHAR", "PARTIAL_MALL"]
DATASET_LABELS = {
    "MHAR": "M_HAR",
    "PARTIAL_MALL": "M_PARTIAL_MALL (IV omitted)",
}


def resolve(path: str | Path) -> Path:
    p = Path(path)
    if not p.is_absolute():
        p = ROOT / p
    return p.resolve()


def atomic_write_json(payload: dict, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")
    tmp.replace(path)


def atomic_write_text(text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(text, encoding="utf-8")
    tmp.replace(path)


def assert_fresh_output_dir(output_dir: Path, allow_existing: bool) -> None:
    if output_dir.exists() and not allow_existing:
        files = [p for p in output_dir.rglob("*") if p.is_file()]
        if files:
            raise SystemExit(
                f"Output directory already contains files: {output_dir}. "
                "Use a fresh directory or pass --allow-existing-output-dir."
            )


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(8)


def set_run_font(run, size: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_cell_borders(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_borders(cell) -> None:
    nil = {"val": "nil"}
    set_cell_borders(cell, top=nil, left=nil, bottom=nil, right=nil, insideH=nil, insideV=nil)


def horizontal_rule(cell, top: bool = False, bottom: bool = False) -> None:
    attrs = {}
    if top:
        attrs["top"] = {"val": "single", "sz": "6", "space": "0", "color": "666666"}
    if bottom:
        attrs["bottom"] = {"val": "single", "sz": "6", "space": "0", "color": "666666"}
    if attrs:
        set_cell_borders(cell, **attrs)


def add_text(paragraph, text: str, size: float = 8.0, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


def add_model_label(paragraph, model: str, size: float = 8.0, bold: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if model in NN_LABELS:
        power, subscript = NN_LABELS[model]
        run = paragraph.add_run("NN")
        set_run_font(run, size, bold=bold)
        sup = paragraph.add_run(str(power))
        set_run_font(sup, size - 1, bold=bold)
        sup.font.superscript = True
        sub = paragraph.add_run(str(subscript))
        set_run_font(sub, size - 1, bold=bold)
        sub.font.subscript = True
        return
    add_text(paragraph, PLAIN_LABELS.get(model, model), size=size, bold=bold)


def add_dataset_label(paragraph, dataset: str, size: float = 8.2, bold: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("M")
    set_run_font(run, size, bold=bold)
    subscript = "HAR" if dataset == "MHAR" else "PARTIAL_MALL"
    sub = paragraph.add_run(subscript)
    set_run_font(sub, size - 0.6, bold=bold)
    sub.font.subscript = True
    if dataset == "PARTIAL_MALL":
        add_text(paragraph, " (IV omitted)", size=size, bold=bold)


def add_caption(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, "Table X ", size=10.5, bold=True)
    add_text(p, "One-day-ahead filtered-historical-simulation VaR results", size=10.5)


def add_note(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, "Notes: ", size=8.2, italic=True)
    note = (
        "We report one-day-ahead 5% filtered-historical-simulation VaR results. "
        "Rel. QL is the cross-sectional average, over stocks, of each model's mean quantile loss "
        "relative to HAR within the same dataset; values below one indicate lower VaR loss than HAR. "
        "Loss sig. reports the percentage of stocks for which a one-sided Diebold-Mariano-style test "
        "on VaR check loss rejects equal loss in favor of the selected model over HAR at the 5% level. "
        "Exc. is the average VaR exceedance rate. UC rej. and IND rej. report the percentage of stocks "
        "for which the Kupiec unconditional-coverage test and Christoffersen independence test reject "
        "at the 5% level. The target exceedance rate is 5%. "
        "M_PARTIAL_MALL is the current replication dataset with IV omitted."
    )
    add_text(p, note, size=8.2)


def load_summary(path: Path, horizon: int) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(path)
    df = pd.read_csv(path)
    required = {
        "dataset",
        "horizon",
        "ticker",
        "model",
        "alpha",
        "exceedance_rate",
        "mean_var_loss",
        "kupiec_p",
        "christoffersen_ind_p",
    }
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Missing columns in VaR summary: {sorted(missing)}")
    df = df[df["horizon"].astype(int).eq(int(horizon)) & df["dataset"].isin(DATASET_ORDER)].copy()
    df["model"] = df["model"].astype(str)
    df["ticker"] = df["ticker"].astype(str).str.upper()
    df["dataset"] = df["dataset"].astype(str)
    df["mean_var_loss"] = pd.to_numeric(df["mean_var_loss"], errors="coerce")
    df["exceedance_rate"] = pd.to_numeric(df["exceedance_rate"], errors="coerce")
    df["kupiec_p"] = pd.to_numeric(df["kupiec_p"], errors="coerce")
    df["christoffersen_ind_p"] = pd.to_numeric(df["christoffersen_ind_p"], errors="coerce")
    df = df.dropna(subset=["mean_var_loss", "exceedance_rate"])
    if df.empty:
        raise ValueError("No VaR summary rows remain after filtering.")
    missing_models = set(MODEL_ORDER) - set(df["model"])
    if missing_models:
        raise ValueError(f"Missing models in VaR summary: {sorted(missing_models)}")
    return df


def compute_loss_significance(var_forecasts_path: Path, horizon: int) -> pd.DataFrame:
    if not var_forecasts_path.exists():
        raise FileNotFoundError(var_forecasts_path)
    usecols = ["date", "ticker", "dataset", "horizon", "model", "var_loss"]
    losses = pd.read_csv(var_forecasts_path, usecols=usecols, parse_dates=["date"])
    losses = losses[
        losses["dataset"].isin(DATASET_ORDER)
        & losses["horizon"].astype(int).eq(int(horizon))
        & losses["model"].isin(MODEL_ORDER)
    ].copy()
    losses["ticker"] = losses["ticker"].astype(str).str.upper()
    losses["model"] = losses["model"].astype(str)
    losses["var_loss"] = pd.to_numeric(losses["var_loss"], errors="coerce")
    losses = losses.dropna(subset=["var_loss"])

    har = losses[losses["model"].eq("HAR")][["date", "ticker", "dataset", "horizon", "var_loss"]].rename(
        columns={"var_loss": "har_var_loss"}
    )
    merged = losses.merge(har, on=["date", "ticker", "dataset", "horizon"], how="inner")

    rows = []
    for (dataset, model, ticker), g in merged.groupby(["dataset", "model", "ticker"], sort=True):
        if model == "HAR":
            dm_stat = float("nan")
            p_value = float("nan")
            sig = False
        else:
            dm_stat, p_value = diebold_mariano(
                g["har_var_loss"].to_numpy(dtype=float),
                g["var_loss"].to_numpy(dtype=float),
                alternative="greater",
                h=int(horizon),
            )
            sig = bool(pd.notna(p_value) and p_value < 0.05)
        rows.append(
            {
                "dataset": dataset,
                "model": model,
                "ticker": ticker,
                "var_loss_dm_stat_vs_har": dm_stat,
                "var_loss_dm_p_vs_har": p_value,
                "var_loss_sig_vs_har_5pct": sig,
                "n_loss_days": int(len(g)),
            }
        )

    per_ticker = pd.DataFrame(rows)
    return (
        per_ticker.groupby(["dataset", "model"], as_index=False)
        .agg(
            loss_sig_share_vs_har_5pct=("var_loss_sig_vs_har_5pct", "mean"),
            loss_test_n_tickers=("ticker", "nunique"),
            loss_test_mean_p_vs_har=("var_loss_dm_p_vs_har", "mean"),
        )
    )


def compute_table(df: pd.DataFrame, loss_sig: pd.DataFrame) -> pd.DataFrame:
    rows = []
    har = (
        df[df["model"].eq("HAR")]
        .set_index(["dataset", "ticker"])["mean_var_loss"]
        .rename("har_loss")
    )
    merged = df.join(har, on=["dataset", "ticker"])
    merged["relative_ql"] = merged["mean_var_loss"] / merged["har_loss"]

    for dataset in DATASET_ORDER:
        for model in MODEL_ORDER:
            g = merged[(merged["dataset"].eq(dataset)) & (merged["model"].eq(model))].copy()
            if g.empty:
                raise ValueError(f"Missing VaR rows for dataset={dataset} model={model}")
            rows.append(
                {
                    "dataset": dataset,
                    "model": model,
                    "n_tickers": int(g["ticker"].nunique()),
                    "relative_quantile_loss": float(g["relative_ql"].mean()),
                    "exceedance_rate": float(g["exceedance_rate"].mean()),
                    "kupiec_reject_share_5pct": float((g["kupiec_p"] < 0.05).mean()),
                    "christoffersen_ind_reject_share_5pct": float((g["christoffersen_ind_p"] < 0.05).mean()),
                    "mean_var_loss": float(g["mean_var_loss"].mean()),
                }
            )
    out = pd.DataFrame(rows)
    out = out.merge(loss_sig, on=["dataset", "model"], how="left")
    if out["loss_sig_share_vs_har_5pct"].isna().any():
        missing = out[out["loss_sig_share_vs_har_5pct"].isna()][["dataset", "model"]]
        raise ValueError(f"Missing VaR loss significance rows: {missing.to_dict(orient='records')[:10]}")
    return out


def build_wide_table(summary: pd.DataFrame) -> pd.DataFrame:
    pieces = []
    for model in MODEL_ORDER:
        row = {"model": model, "label": PLAIN_LABELS.get(model, model)}
        for dataset in DATASET_ORDER:
            g = summary[(summary["dataset"].eq(dataset)) & (summary["model"].eq(model))]
            if g.empty:
                raise ValueError(f"Missing summary row dataset={dataset} model={model}")
            rec = g.iloc[0]
            prefix = dataset.lower()
            row[f"{prefix}_relative_quantile_loss"] = rec["relative_quantile_loss"]
            row[f"{prefix}_loss_sig_share_vs_har_5pct"] = rec["loss_sig_share_vs_har_5pct"]
            row[f"{prefix}_exceedance_rate"] = rec["exceedance_rate"]
            row[f"{prefix}_kupiec_reject_share_5pct"] = rec["kupiec_reject_share_5pct"]
            row[f"{prefix}_christoffersen_ind_reject_share_5pct"] = rec["christoffersen_ind_reject_share_5pct"]
            row[f"{prefix}_n_tickers"] = int(rec["n_tickers"])
        pieces.append(row)
    return pd.DataFrame(pieces)


def fmt_rel(value: float) -> str:
    return f"{value:.3f}"


def fmt_pct(value: float) -> str:
    return f"{100.0 * value:.1f}"


def build_docx(wide: pd.DataFrame, output_path: Path) -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_caption(doc)

    table = doc.add_table(rows=2 + len(wide), cols=1 + 5 * len(DATASET_ORDER))
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            clear_borders(cell)

    table.cell(0, 0).merge(table.cell(1, 0))
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Model", size=8.2, bold=True)

    start = 1
    for dataset in DATASET_ORDER:
        merged = table.cell(0, start).merge(table.cell(0, start + 4))
        add_dataset_label(merged.paragraphs[0], dataset, size=8.2, bold=True)
        for j, label in enumerate(["Rel. QL", "Loss sig.", "Exc.", "UC rej.", "IND rej."]):
            p = table.cell(1, start + j).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, label, size=8.0, bold=True)
        start += 5

    best_rel = {}
    for dataset in DATASET_ORDER:
        col = f"{dataset.lower()}_relative_quantile_loss"
        best_rel[dataset] = float(wide[col].min())

    for i, (_, row) in enumerate(wide.iterrows(), start=2):
        add_model_label(table.cell(i, 0).paragraphs[0], row["model"], size=8.0)
        col_idx = 1
        for dataset in DATASET_ORDER:
            prefix = dataset.lower()
            values = [
                (fmt_rel(row[f"{prefix}_relative_quantile_loss"]), row[f"{prefix}_relative_quantile_loss"]),
                (fmt_pct(row[f"{prefix}_loss_sig_share_vs_har_5pct"]), None),
                (fmt_pct(row[f"{prefix}_exceedance_rate"]), None),
                (fmt_pct(row[f"{prefix}_kupiec_reject_share_5pct"]), None),
                (fmt_pct(row[f"{prefix}_christoffersen_ind_reject_share_5pct"]), None),
            ]
            for text, raw in values:
                p = table.cell(i, col_idx).paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                bold = raw is not None and abs(float(raw) - best_rel[dataset]) < 1e-12
                add_text(p, text, size=8.0, bold=bold)
                col_idx += 1

    for cell in table.rows[0].cells:
        horizontal_rule(cell, top=True)
    for cell in table.rows[1].cells:
        horizontal_rule(cell, bottom=True)
    for cell in table.rows[-1].cells:
        horizontal_rule(cell, bottom=True)

    add_note(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def write_markdown(wide: pd.DataFrame, path: Path) -> None:
    headers = [
        "Model",
        "M_HAR Rel. QL",
        "M_HAR Loss sig.",
        "M_HAR Exc.",
        "M_HAR UC rej.",
        "M_HAR IND rej.",
        "M_PARTIAL_MALL Rel. QL",
        "M_PARTIAL_MALL Loss sig.",
        "M_PARTIAL_MALL Exc.",
        "M_PARTIAL_MALL UC rej.",
        "M_PARTIAL_MALL IND rej.",
    ]
    lines = [
        "# Table X. One-day-ahead filtered-historical-simulation VaR results",
        "",
        "|" + "|".join(headers) + "|",
        "|" + "|".join(["---"] * len(headers)) + "|",
    ]
    for _, row in wide.iterrows():
        model_label = PLAIN_LABELS.get(row["model"], row["model"])
        if row["model"] in NN_LABELS:
            power, subscript = NN_LABELS[row["model"]]
            model_label = f"NN_{subscript}^{power}"
        vals = [
            model_label,
            fmt_rel(row["mhar_relative_quantile_loss"]),
            fmt_pct(row["mhar_loss_sig_share_vs_har_5pct"]),
            fmt_pct(row["mhar_exceedance_rate"]),
            fmt_pct(row["mhar_kupiec_reject_share_5pct"]),
            fmt_pct(row["mhar_christoffersen_ind_reject_share_5pct"]),
            fmt_rel(row["partial_mall_relative_quantile_loss"]),
            fmt_pct(row["partial_mall_loss_sig_share_vs_har_5pct"]),
            fmt_pct(row["partial_mall_exceedance_rate"]),
            fmt_pct(row["partial_mall_kupiec_reject_share_5pct"]),
            fmt_pct(row["partial_mall_christoffersen_ind_reject_share_5pct"]),
        ]
        lines.append("|" + "|".join(vals) + "|")
    lines.extend(
        [
            "",
            "Notes: Rel. QL is mean VaR quantile loss relative to HAR within the same dataset. "
            "Loss sig. reports the percentage of stocks for which a one-sided DM-style VaR check-loss "
            "test favors the selected model over HAR at the 5% level. Exc., UC rej., and IND rej. are "
            "reported in percent. The VaR tail probability is 5%.",
        ]
    )
    atomic_write_text("\n".join(lines) + "\n", path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a Word/CSV table for FHS VaR summary results.")
    parser.add_argument("--input", default=str(DEFAULT_INPUT))
    parser.add_argument("--var-forecasts", default=str(DEFAULT_VAR_FORECASTS))
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR))
    parser.add_argument("--horizon", type=int, default=1)
    parser.add_argument("--allow-existing-output-dir", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = resolve(args.input)
    var_forecasts_path = resolve(args.var_forecasts)
    output_dir = resolve(args.output_dir)
    assert_fresh_output_dir(output_dir, args.allow_existing_output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    raw = load_summary(input_path, horizon=int(args.horizon))
    loss_sig = compute_loss_significance(var_forecasts_path, horizon=int(args.horizon))
    summary = compute_table(raw, loss_sig=loss_sig)
    wide = build_wide_table(summary)

    long_path = output_dir / "var_summary_table_long.csv"
    wide_path = output_dir / "var_summary_table_wide.csv"
    docx_path = output_dir / "var_summary_table_h1.docx"
    md_path = output_dir / "var_summary_table_h1.md"
    summary.to_csv(long_path, index=False)
    wide.to_csv(wide_path, index=False)
    build_docx(wide, docx_path)
    write_markdown(wide, md_path)

    atomic_write_json(
        {
            "created_utc": datetime.now(timezone.utc).isoformat(),
            "source": str(input_path),
            "var_forecasts_source": str(var_forecasts_path),
            "output_dir": str(output_dir),
            "horizon": int(args.horizon),
            "datasets": DATASET_ORDER,
            "models": MODEL_ORDER,
            "metrics": {
                "relative_quantile_loss": "ticker-level mean_var_loss divided by HAR mean_var_loss, averaged across tickers",
                "loss_sig_share_vs_har_5pct": "share of tickers where one-sided DM-style VaR check-loss test favors the model over HAR at 5%",
                "exceedance_rate": "cross-sectional average exceedance rate",
                "kupiec_reject_share_5pct": "share of tickers with Kupiec p < 0.05",
                "christoffersen_ind_reject_share_5pct": "share of tickers with Christoffersen independence p < 0.05",
            },
            "outputs": {
                "long_csv": str(long_path),
                "wide_csv": str(wide_path),
                "docx": str(docx_path),
                "markdown": str(md_path),
            },
        },
        output_dir / "run_provenance.json",
    )

    print(f"Wrote {docx_path}")
    print(f"Wrote {wide_path}")
    print(f"Wrote {md_path}")


if __name__ == "__main__":
    main()
