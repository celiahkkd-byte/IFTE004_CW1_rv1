from __future__ import annotations

import argparse
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TRADING_DAYS_PER_YEAR = 252


def _rv_to_annualized_vol_percent(s: pd.Series) -> pd.Series:
    return np.sqrt(pd.to_numeric(s, errors='coerce').clip(lower=0.0) * TRADING_DAYS_PER_YEAR) * 100.0


VARIABLES: list[dict] = [
    {
        'no': 1,
        'acronym': 'RVD',
        'source_col': 'rvd',
        'asset_specific': True,
        'transform': _rv_to_annualized_vol_percent,
        'available_if_missing': False,
    },
    {
        'no': 2,
        'acronym': 'RVW',
        'source_col': 'rvw',
        'asset_specific': True,
        'transform': _rv_to_annualized_vol_percent,
        'available_if_missing': False,
    },
    {
        'no': 3,
        'acronym': 'RVM',
        'source_col': 'rvm',
        'asset_specific': True,
        'transform': _rv_to_annualized_vol_percent,
        'available_if_missing': False,
    },
    {
        'no': 4,
        'acronym': 'IV',
        'source_col': 'iv',
        'asset_specific': True,
        'transform': lambda s: pd.to_numeric(s, errors='coerce'),
        'available_if_missing': True,
        'missing_reason': 'omitted: firm-level OptionMetrics IV not supplied',
    },
    {
        'no': 5,
        'acronym': 'EA',
        'source_col': 'ea',
        'asset_specific': True,
        'transform': lambda s: pd.to_numeric(s, errors='coerce'),
        'omit_stats': True,
        'available_if_missing': False,
    },
    {
        'no': 6,
        'acronym': 'VIX',
        'source_col': 'vix',
        'asset_specific': False,
        'transform': lambda s: pd.to_numeric(s, errors='coerce'),
        'available_if_missing': False,
    },
    {
        'no': 7,
        'acronym': 'EPU',
        'source_col': 'epu',
        'asset_specific': False,
        'transform': lambda s: pd.to_numeric(s, errors='coerce'),
        'available_if_missing': False,
    },
    {
        'no': 8,
        'acronym': 'US3M',
        'source_col': 'us3m_diff',
        'asset_specific': False,
        'transform': lambda s: pd.to_numeric(s, errors='coerce') * 100.0,
        'available_if_missing': False,
    },
    {
        'no': 9,
        'acronym': 'HSI',
        'source_col': 'hsi',
        'asset_specific': False,
        'transform': lambda s: pd.to_numeric(s, errors='coerce') * 100.0,
        'available_if_missing': False,
    },
    {
        'no': 10,
        'acronym': 'M1W',
        'source_col': 'm1w',
        'asset_specific': True,
        'transform': lambda s: pd.to_numeric(s, errors='coerce') * 100.0,
        'available_if_missing': False,
    },
    {
        'no': 11,
        'acronym': '$VOL',
        'source_col': 'dvol',
        'asset_specific': True,
        'transform': lambda s: pd.to_numeric(s, errors='coerce') * 100.0,
        'available_if_missing': False,
    },
    {
        'no': 12,
        'acronym': 'ADS',
        'source_col': 'ads',
        'asset_specific': False,
        'transform': lambda s: pd.to_numeric(s, errors='coerce') * 100.0,
        'available_if_missing': False,
    },
]


STAT_SPECS: list[tuple[str, str, Callable[[pd.Series], float]]] = [
    ('mean', 'Mean', lambda x: float(x.mean())),
    ('median', 'Median', lambda x: float(x.median())),
    ('maximum', 'Maximum', lambda x: float(x.max())),
    ('minimum', 'Minimum', lambda x: float(x.min())),
    ('std', 'Standard deviation', lambda x: float(x.std(ddof=1))),
    ('skewness', 'Skewness', lambda x: float(x.skew())),
    ('kurtosis', 'Kurtosis', lambda x: float(x.kurt() + 3.0)),
]


def _resolve(path: str | Path) -> Path:
    return Path(path).expanduser().resolve()


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _format_number(value: float | None) -> str:
    if value is None or not np.isfinite(value):
        return ''
    rounded = round(float(value), 2)
    if rounded == -0.0:
        rounded = 0.0
    return f'{rounded:.2f}'


def _stat_values(x: pd.Series) -> dict[str, float]:
    clean = pd.to_numeric(x, errors='coerce').dropna()
    if clean.empty:
        return {key: np.nan for key, _, _ in STAT_SPECS}
    return {key: fn(clean) for key, _, fn in STAT_SPECS}


def _asset_ranges(panel: pd.DataFrame, values: pd.Series) -> dict[str, tuple[float, float]]:
    tmp = pd.DataFrame({'ticker': panel['ticker'].astype(str), 'value': values}).dropna(subset=['value'])
    ranges: dict[str, tuple[float, float]] = {}
    if tmp.empty:
        return {key: (np.nan, np.nan) for key, _, _ in STAT_SPECS}
    for key, _, fn in STAT_SPECS:
        by_ticker = tmp.groupby('ticker')['value'].apply(lambda s: fn(s.dropna()))
        ranges[key] = (float(by_ticker.min()), float(by_ticker.max()))
    return ranges


def _cell_text(value: float, interval: tuple[float, float] | None, show_interval: bool) -> str:
    base = _format_number(value)
    if not show_interval or interval is None:
        return base
    lo, hi = interval
    return f'{base} [{_format_number(lo)},{_format_number(hi)}]'


def compute_summary(panel: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    if 'ticker' not in panel.columns:
        raise ValueError('forecasting panel must contain ticker column')

    table_rows: list[dict] = []
    audit_rows: list[dict] = []

    for spec in VARIABLES:
        row: dict[str, str | int] = {
            'No.': int(spec['no']),
            'Acronym': str(spec['acronym']),
        }
        source_col = str(spec['source_col'])
        if source_col not in panel.columns:
            if spec.get('available_if_missing'):
                for key, label, _ in STAT_SPECS:
                    row[label] = ''
                    audit_rows.append(
                        {
                            'acronym': spec['acronym'],
                            'source_col': source_col,
                            'stat': key,
                            'overall': np.nan,
                            'asset_min': np.nan,
                            'asset_max': np.nan,
                            'note': spec.get('missing_reason', 'missing'),
                        }
                    )
                table_rows.append(row)
                continue
            raise ValueError(f'Missing required source column for {spec["acronym"]}: {source_col}')

        values = spec['transform'](panel[source_col])
        omit_stats = bool(spec.get('omit_stats', False))
        overall = _stat_values(values)
        ranges = _asset_ranges(panel, values) if bool(spec.get('asset_specific')) else {}

        for key, label, _ in STAT_SPECS:
            interval = ranges.get(key)
            audit_rows.append(
                {
                    'acronym': spec['acronym'],
                    'source_col': source_col,
                    'stat': key,
                    'overall': overall[key],
                    'asset_min': interval[0] if interval is not None else np.nan,
                    'asset_max': interval[1] if interval is not None else np.nan,
                    'note': 'stats omitted for binary EA dummy' if omit_stats else '',
                }
            )
            if omit_stats:
                row[label] = ''
            else:
                row[label] = _cell_text(overall[key], interval, show_interval=interval is not None)
        table_rows.append(row)

    return pd.DataFrame(table_rows), pd.DataFrame(audit_rows)


def _set_run_font(run, size: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _clear_cell(cell) -> None:
    for p in cell.paragraphs:
        p.clear()


def _set_cell_text(cell, text: str, size: float = 7.4, bold: bool = False, italic: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    _set_run_font(run, size=size, bold=bold, italic=italic)


def _set_cell_borders(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        element = tc_borders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f'w:{key}'), str(value))


def _clear_borders(cell) -> None:
    nil = {'val': 'nil'}
    _set_cell_borders(cell, top=nil, left=nil, bottom=nil, right=nil, insideH=nil, insideV=nil)


def _horizontal_rule(cell, top: bool = False, bottom: bool = False) -> None:
    attrs = {}
    if top:
        attrs['top'] = {'val': 'single', 'sz': '6', 'space': '0', 'color': '666666'}
    if bottom:
        attrs['bottom'] = {'val': 'single', 'sz': '6', 'space': '0', 'color': '666666'}
    if attrs:
        _set_cell_borders(cell, **attrs)


def _set_cell_width(cell, width: float) -> None:
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in('w:tcW')
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(int(width * 1440)))
    tc_w.set(qn('w:type'), 'dxa')


def write_word_table(table_df: pd.DataFrame, output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(8)

    caption = doc.add_paragraph()
    caption.paragraph_format.space_after = Pt(4)
    r = caption.add_run('Table 1 ')
    _set_run_font(r, 10.5, bold=True)
    r = caption.add_run('List and summary statistics of explanatory variables')
    _set_run_font(r, 10.5)

    cols = table_df.columns.tolist()
    table = doc.add_table(rows=len(table_df) + 1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.32, 0.70, 1.32, 1.32, 1.32, 1.32, 1.55, 1.30, 1.30]

    for row in table.rows:
        for j, cell in enumerate(row.cells):
            _clear_borders(cell)
            _set_cell_width(cell, widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for j, col in enumerate(cols):
        _set_cell_text(table.cell(0, j), col, size=7.8, bold=False)
    for cell in table.rows[0].cells:
        _horizontal_rule(cell, top=True, bottom=True)

    for i, (_, row) in enumerate(table_df.iterrows(), start=1):
        for j, col in enumerate(cols):
            align = WD_ALIGN_PARAGRAPH.LEFT if col == 'Acronym' else WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(table.cell(i, j), row[col], size=7.1, align=align)

    for cell in table.rows[-1].cells:
        _horizontal_rule(cell, bottom=True)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(5)
    r = note.add_run('Notes: ')
    _set_run_font(r, 8.2, italic=True)
    text = (
        'EA is equal to one if the company has an earnings announcement on the day of the forecast, zero otherwise; '
        'therefore, descriptive statistics are omitted for EA. Square brackets contain interval-based measures for '
        'asset-specific variables, defined as the minimum and maximum value of the statistic when calculated '
        'individually across tickers. RVD, RVW, and RVM are reported as annualized volatility percentages computed '
        'from realized variance. US3M and $VOL are reported for the transformed variables, and ADS is reported in '
        'percent. Firm-level OptionMetrics IV is unavailable in this reproduction and is omitted from PARTIAL_MALL.'
    )
    r = note.add_run(text)
    _set_run_font(r, 8.2)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description='Create paper-style Table 1 explanatory-variable summary statistics.')
    parser.add_argument('--panel', default='data/processed/forecasting_panel.csv')
    parser.add_argument('--output-dir', required=True)
    parser.add_argument('--docx-name', default='table1_explanatory_variable_summary.docx')
    args = parser.parse_args()

    panel_path = _resolve(args.panel)
    output_dir = _resolve(args.output_dir)
    if not panel_path.exists():
        raise FileNotFoundError(f'Missing panel file: {panel_path}')

    panel = pd.read_csv(panel_path, parse_dates=['date'])
    table_df, audit_df = compute_summary(panel)

    output_dir.mkdir(parents=True, exist_ok=True)
    csv_path = output_dir / 'table1_explanatory_variable_summary.csv'
    audit_path = output_dir / 'table1_explanatory_variable_summary_audit.csv'
    docx_path = output_dir / args.docx_name
    provenance_path = output_dir / 'run_provenance.json'

    table_df.to_csv(csv_path, index=False)
    audit_df.to_csv(audit_path, index=False)
    write_word_table(table_df, docx_path)
    provenance_path.write_text(
        json.dumps(
            {
                'created_at_utc': _utc_now(),
                'panel': str(panel_path),
                'rows': int(len(panel)),
                'tickers': int(panel['ticker'].nunique()),
                'date_min': str(panel['date'].min().date()) if 'date' in panel else None,
                'date_max': str(panel['date'].max().date()) if 'date' in panel else None,
                'outputs': {
                    'docx': str(docx_path),
                    'csv': str(csv_path),
                    'audit_csv': str(audit_path),
                },
            },
            indent=2,
            sort_keys=True,
        ),
        encoding='utf-8',
    )

    print(f'Wrote {docx_path}')
    print(f'Wrote {csv_path}')
    print(f'Wrote {audit_path}')


if __name__ == '__main__':
    main()
