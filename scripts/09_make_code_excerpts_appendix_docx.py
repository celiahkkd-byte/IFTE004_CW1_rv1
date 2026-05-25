"""Create a report appendix with relevant reproduction code excerpts.

The output is a copyable Word/Markdown appendix. It documents the key code paths
used for data construction, preprocessing, estimation, NN selection, and
evaluation. It does not execute any model code.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_ROOT / "outputs_code_excerpts_appendix_20260524"
DOCX_PATH = OUT_DIR / "appendix_code_excerpts_reproduction.docx"
MD_PATH = OUT_DIR / "appendix_code_excerpts_reproduction.md"
MAINLINE_DIR = "outputs_final_core_with_bagging_gb_harfix_nn50_tuned_no_refit_h1h5_20260523"


def excerpts() -> list[dict[str, str]]:
    return [
        {
            "section": "Raw intraday data to realized measures",
            "file": "src/rv1rep/intraday.py::compute_daily_realized_measures",
            "excerpt": """r5 = _daily_5min_returns(
    g,
    expected_bars=expected_bars,
    bar_interval_minutes=bar_interval_minutes,
)
rv = float(np.sum(r5 ** 2) * rv_scale)
rvp = float(np.sum(r5[r5 > 0] ** 2) * rv_scale)
rvn = float(np.sum(r5[r5 < 0] ** 2) * rv_scale)
rq = float((n / 3.0) * np.sum(r5 ** 4) * (rv_scale ** 2))
dollar_volume = float((g['close'] * g['volume']).sum())""",
            "relevance": "Shows how the replication constructs daily realized variance, semivariances, realized quarticity, and dollar volume from five-minute intraday bars.",
        },
        {
            "section": "Forecast target alignment",
            "file": "src/rv1rep/features.py::_future_target and add_targets",
            "excerpt": """def _future_target(s: pd.Series, horizon: int, mode: str) -> pd.Series:
    if horizon == 1:
        return s.shift(-1)
    if mode == 'future_average':
        return (
            s.shift(-1)
             .rolling(horizon, min_periods=horizon)
             .mean()
             .shift(-(horizon - 1))
        )

out[f'target_rv_h{h}'] = out.groupby('ticker')['rv'].transform(
    lambda x: _future_target(x, h, mode)
)""",
            "relevance": "Documents that one-day-ahead and multi-day targets are future values, aligned at date t without using future information in the predictors.",
        },
        {
            "section": "Lagged HAR and extended predictors",
            "file": "src/rv1rep/features.py::add_asset_features",
            "excerpt": """df['rvd'] = df['rv']
df['rvw'] = g['rv'].transform(
    lambda x: x.rolling(weekly_window, min_periods=weekly_window).mean()
)
df['rvm'] = g['rv'].transform(
    lambda x: x.rolling(monthly_window, min_periods=monthly_window).mean()
)
df['m1w'] = g['cc_logret'].transform(
    lambda x: x.rolling(weekly_window, min_periods=weekly_window).sum()
)
df['dvol'] = g['dollar_volume'].transform(
    lambda x: np.log(np.maximum(x, eps)).diff()
)""",
            "relevance": "Records the construction of the standard HAR variables and selected MALL-style public predictors used in the extended dataset.",
        },
        {
            "section": "External variables and IV omission",
            "file": "src/rv1rep/features.py::feature_columns_for_model",
            "excerpt": """if dataset.upper() == 'MHAR':
    return [c for c in base if c in cols]

if model == 'LOGHAR':
    extra_candidates = [
        'log_iv', 'ea', 'm1w', 'dvol', 'log_vix',
        'hsi', 'ads', 'us3m_diff', 'epu'
    ]
else:
    extra_candidates = [
        'iv', 'ea', 'm1w', 'dvol', 'vix',
        'hsi', 'ads', 'us3m_diff', 'epu'
    ]
extras = [c for c in extra_candidates if c in cols]""",
            "relevance": "Shows the implemented distinction between MHAR and the public extended dataset. IV is included only if the relevant panel column exists; otherwise the replication is PARTIAL_MALL, i.e. IV omitted.",
        },
        {
            "section": "Feature standardization",
            "file": "config/paper_core_rolling_tuned_no_refit.yaml; src/rv1rep/preprocessing.py::Standardizer",
            "excerpt": """preprocessing:
  standardize_binary_features: true
  binary_features: [ea]

def standardizer_from_config(cfg: dict) -> Standardizer:
    prep = cfg.get('preprocessing', {})
    return Standardizer(
        categorical_features=tuple(prep.get('binary_features', CATEGORICAL_FEATURES)),
        standardize_binary_features=bool(prep.get('standardize_binary_features', True)),
    )

if self.standardize_binary_features:
    self.categorical_columns_ = []
    self.continuous_columns_ = list(self.columns_)""",
            "relevance": "Documents the mainline feature standardization policy: all feature columns supplied to a model are standardized within the in-sample training window, including the EA dummy when it is present in PARTIAL_MALL. Robustness scripts can switch this policy off explicitly for binary EA.",
        },
        {
            "section": "Rolling estimation and refit frequency",
            "file": "config/paper_core_rolling_tuned_no_refit.yaml; src/rv1rep/forecasting.py",
            "excerpt": """# Mainline configuration
ml_refit_every: 5

train_val_window = train_n + val_n
non_tuned_models = {
    'HAR', 'HARX', 'LOGHAR', 'LEVHAR', 'SHAR', 'HARQ',
    'BAGGING', 'RANDOMFOREST'
}
daily_refit_models = {
    'HAR', 'HARX', 'LOGHAR', 'LEVHAR', 'SHAR', 'HARQ',
    'RIDGE', 'LASSO', 'ELASTICNET', 'ADAPTIVELASSO', 'POSTLASSO'
}
refit = (
    model_u in daily_refit_models
    or last_est is None
    or (i - last_refit_i >= refit_every)
)""",
            "relevance": "Shows the rolling-window implementation. HAR-family and regularized linear models are refitted daily. For computational tractability, Bagging, RandomForest, and GradientBoosting are refitted every five test observations in the mainline, which is a deviation from the paper's daily rolling refit interpretation.",
        },
        {
            "section": "Train/validation split inside the rolling window",
            "file": "src/rv1rep/forecasting.py::_fit_one_asset_rolling",
            "excerpt": """if model_u in non_tuned_models:
    # Non-tuned rolling models use the whole train+validation window
    # as one in-sample fit block.
    train_dates = window_dates
    val_dates = window_dates[:0]
else:
    train_dates = window_dates[:train_n]
    val_dates = window_dates[train_n:]

scaler = Standardizer().fit(train[feature_cols])
X_train = scaler.transform(train[feature_cols])
X_val = scaler.transform(val[feature_cols])
est, params = fit_sklearn_model(
    model_name, X_train, y_train, X_val, y_val, cfg,
)""",
            "relevance": "Documents the paper-style estimation split: non-tuned models use the combined in-sample window, while tuned models reserve a validation block for hyperparameter selection.",
        },
        {
            "section": "Validation tuning without train-validation refit",
            "file": "src/rv1rep/models.py::_maybe_refit_selected_tuned_model",
            "excerpt": """def _maybe_refit_selected_tuned_model(
    est, X_train, y_train, X_val, y_val, cfg
):
    if _refit_tuned_models_on_train_validation(cfg):
        est.fit(pd.concat([X_train, X_val]), pd.concat([y_train, y_val]))
        return est, 'train_plus_validation_after_selection'
    return est, 'train_only_after_validation_selection'

key, est, loss = select_by_validation(
    candidates, X_train, y_train, X_val, y_val
)""",
            "relevance": "Provides the central evidence that validation-tuned models select hyperparameters on validation data and then keep the selected training-fitted estimator, matching the no-refit design used in the final mainline.",
        },
        {
            "section": "Tree and gradient boosting configuration",
            "file": "src/rv1rep/models.py::fit_sklearn_model",
            "excerpt": """BaggingRegressor(
    estimator=DecisionTreeRegressor(
        min_samples_leaf=tree_cfg['min_samples_leaf']
    ),
    n_estimators=tree_cfg['n_estimators'],
)

RandomForestRegressor(
    n_estimators=tree_cfg['n_estimators'],
    min_samples_leaf=tree_cfg['min_samples_leaf'],
    max_features=max_features,
)

chains = [
    (int(depth), float(lr))
    for depth in gb_cfg['depths']
    for lr in gb_cfg['learning_rates']
]""",
            "relevance": "Shows the implementation of the paper-style tree settings and the validation grid for Gradient Boosting.",
        },
        {
            "section": "Neural-network architecture and training",
            "file": "config/paper_core_rolling_tuned_no_refit.yaml; src/rv1rep/nn.py",
            "excerpt": """# Mainline configuration
neural_network:
  seeds: 50
  ensemble_top: 10
  epochs: 500
  patience: 100
  learning_rate: 0.001
  dropout: 0.8

for units in hidden_layers:
    x = tf.keras.layers.Dense(
        units, kernel_initializer='glorot_normal'
    )(x)
    x = tf.keras.layers.LeakyReLU(alpha=0.01)(x)
    x = tf.keras.layers.Dropout(dropout)(x)

opt = tf.keras.optimizers.Adam(learning_rate=learning_rate)
model.compile(optimizer=opt, loss='mse')""",
            "relevance": "Documents the NN ingredients reported in the hyperparameter table. The mainline config explicitly sets seeds=50 and ensemble_top=10, which is also recorded in the final prediction params.",
        },
        {
            "section": "NN checkpoint aggregation and top-10 ensemble",
            "file": "scripts/04_run_nn_checkpoints.py::aggregate_nn_checkpoints",
            "excerpt": """available.append((seed, float(df_seed['val_mse'].iloc[0]), path))
if require_all_seeds and len(available) != len(seeds):
    raise RuntimeError('Missing seed checkpoints')

selected = sorted(available, key=lambda x: x[1])[:ensemble_top]
seed_frames = [_load_seed_file(path) for _, _, path in selected]
raw_matrix = np.vstack([
    f['forecast_raw'].to_numpy(dtype=float) for f in seed_frames
])
raw_pred = raw_matrix.mean(axis=0)""",
            "relevance": "Shows how the checkpointed NN run enforces complete seed availability and constructs the top-10 ensemble by validation MSE.",
        },
        {
            "section": "Forecast evaluation and pairwise DM tests",
            "file": "src/rv1rep/evaluation.py::forecast_metrics and pairwise_relative_mse",
            "excerpt": """e = g['actual_rv'].to_numpy() - g['forecast_rv'].to_numpy()
mse = float(np.mean(e ** 2))
out['relative_mse_vs_har'] = out['mse'] / out['har_mse']

li = (merged['actual_rv'] - merged['f_i']) ** 2
lj = (merged['actual_rv'] - merged['f_j']) ** 2
mse_i, mse_j = float(li.mean()), float(lj.mean())
row[col_model] = float(np.mean(ratios))
stat, p = diebold_mariano(li.to_numpy(), lj.to_numpy(), h=int(horizon))""",
            "relevance": "Documents how out-of-sample MSE, relative MSE, and one-sided Diebold-Mariano tests are produced from final test-period forecasts.",
        },
    ]


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_text(cell, text: str, *, bold: bool = False, size: int = 7, mono: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Courier New" if mono else "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_paragraph(document: Document, text: str, *, bold: bool = False, size: int = 10) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def write_markdown(rows: list[dict[str, str]]) -> None:
    lines = [
        "# Appendix B.X Code excerpts for the reproduction",
        "",
        "The table below reports short excerpts from the most relevant parts of the replication code. The excerpts are abridged for readability; file and function names identify the full implementation.",
        "",
        f"Scope: this appendix describes the reported mainline results in `{MAINLINE_DIR}/`, with final predictions in `{MAINLINE_DIR}/predictions/model_predictions.csv`. This mainline contains h=1 and h=5 for MHAR and PARTIAL_MALL, includes GradientBoosting and the NN50 results, and uses the validation-tuned no-refit Regularized/GradientBoosting rows.",
        "",
        "| Part | Source file / function | Code excerpt | Relevance |",
        "|---|---|---|---|",
    ]
    for row in rows:
        code = "<br>".join(row["excerpt"].splitlines()).replace("|", "\\|")
        lines.append(
            f"| {row['section']} | `{row['file']}` | `{code}` | {row['relevance']} |"
        )
    lines.append("")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx(rows: list[dict[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Appendix B.X Code Excerpts for the Reproduction")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    add_paragraph(
        document,
        "The table below reports short excerpts from the most relevant parts of the replication code. "
        "The excerpts are abridged for readability; file and function names identify the full implementation.",
    )
    add_paragraph(
        document,
        f"Scope: this appendix describes the reported mainline results in {MAINLINE_DIR}/, with final predictions in "
        f"{MAINLINE_DIR}/predictions/model_predictions.csv. This mainline contains h=1 and h=5 for MHAR and "
        "PARTIAL_MALL, includes GradientBoosting and the NN50 results, and uses the validation-tuned no-refit "
        "Regularized/GradientBoosting rows.",
    )

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Part", "Source file / function", "Code excerpt", "Why this is relevant"]
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "D9EAF7")
        set_text(cell, header, bold=True, size=7)

    widths = [Cm(3.3), Cm(5.1), Cm(11.6), Cm(7.4)]
    for row in rows:
        cells = table.add_row().cells
        values = [row["section"], row["file"], row["excerpt"], row["relevance"]]
        for i, (cell, text, width) in enumerate(zip(cells, values, widths)):
            cell.width = width
            set_text(cell, text, size=6 if i == 2 else 7, mono=(i == 2))

    document.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = excerpts()
    write_markdown(rows)
    build_docx(rows)
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
